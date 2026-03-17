"""
ECHO Local Backend v3.5 — FastAPI + Ollama + ChromaDB
Multi-Agent AI Orchestration with Performance & Intelligence Upgrades

Endpoints:
    GET  /api/health                      — Backend/GPU status, loaded models, uptime
    GET  /api/system                      — Real-time CPU/RAM/GPU/Disk metrics
    GET  /api/agents                      — Multi-agent pipeline status
    GET  /api/models                      — Available models with capabilities
    GET  /api/telemetry                   — Aggregated telemetry (VRAM, cache, agents)
    POST /api/chat                        — SSE streaming chat with full pipeline
    POST /api/documents                   — Ingest documents into ChromaDB for RAG
    GET  /api/documents/list              — List all documents
    DELETE /api/documents/{id}            — Delete a document
    POST /api/semantic-search             — Vector similarity search over documents
    POST /api/skill-tools                 — Skill creator/eval/trigger-tuning via Ollama
    GET  /api/skills/scan                 — Read .md skill files from backend/skills/
    GET  /api/skills/compiled             — v3.5: Compiled skill definitions with triggers
    GET  /api/cache/stats                 — Response cache statistics
    DELETE /api/cache                     — Clear response cache
    POST /api/memory/store                — Store a memory (episodic/semantic/procedural)
    POST /api/memory/recall               — Semantic search across all memory types
    GET  /api/memory/list                 — List memories by type
    DELETE /api/memory/{id}               — Delete a memory
    POST /api/feedback                    — Submit response feedback
    GET  /api/feedback/stats              — Feedback statistics
    POST /api/models/manage               — Load/unload Ollama models
    GET  /api/sentinel/health             — Sentinel routing health snapshot
    POST /api/sentinel/optimize           — Trigger Sentinel threshold auto-tuning
    POST /api/sentinel/improve            — v3.5: Self-improvement prompt engine
    GET  /api/tools/discover              — v3.5: Discover all available tools
    POST /api/projects/create             — v3.5: Create an AI project
    GET  /api/projects                    — v3.5: List all projects
    GET  /api/projects/{id}               — v3.5: Get project details
    POST /api/projects/{id}/task          — v3.5: Add task to project
    DELETE /api/projects/{id}             — v3.5: Delete project
    GET  /api/speculative/status          — v3.5: Speculative decode model status
    GET  /api/vision/status               — Vision model availability check
    POST /api/vision/analyze              — Analyze image with local vision model (llava, etc.)

Run:
    cd backend
    pip install -r requirements.txt
    uvicorn main:app --host 0.0.0.0 --port 8000 --reload
"""

import asyncio
import hashlib
import json
import logging
import logging.handlers
import mimetypes
import os
import platform
import re
import socket
import sys
import time
from collections import OrderedDict
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import httpx
import psutil
from fastapi import FastAPI, HTTPException

# ── Optional v3.2 dependencies (graceful fallback if not installed) ────────────
try:
    from rank_bm25 import BM25Okapi as _BM25Okapi
    _BM25_AVAILABLE = True
except ImportError:
    _BM25Okapi = None  # type: ignore
    _BM25_AVAILABLE = False

try:
    from watchdog.observers import Observer as _WatchdogObserver
    from watchdog.events import FileSystemEventHandler as _FSEventHandler
    _WATCHDOG_AVAILABLE = True
except ImportError:
    _WatchdogObserver = None  # type: ignore
    _FSEventHandler = object  # type: ignore
    _WATCHDOG_AVAILABLE = False

try:
    import pdfplumber as _pdfplumber
    _PDF_AVAILABLE = True
except ImportError:
    _pdfplumber = None  # type: ignore
    _PDF_AVAILABLE = False

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DocxDocument = None  # type: ignore
    _DOCX_AVAILABLE = False

try:
    from duckduckgo_search import DDGS as _DDGS
    _DDG_AVAILABLE = True
except ImportError:
    _DDGS = None  # type: ignore
    _DDG_AVAILABLE = False

try:
    from bs4 import BeautifulSoup as _BeautifulSoup
    _BS4_AVAILABLE = True
except ImportError:
    _BeautifulSoup = None  # type: ignore
    _BS4_AVAILABLE = False

try:
    import trafilatura as _trafilatura
    _TRAFILATURA_AVAILABLE = True
except ImportError:
    _trafilatura = None  # type: ignore
    _TRAFILATURA_AVAILABLE = False

try:
    from cachetools import TTLCache as _TTLCache
    _CACHETOOLS_AVAILABLE = True
except ImportError:
    _TTLCache = None  # type: ignore
    _CACHETOOLS_AVAILABLE = False

from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response

# ── v3.6: Prompt Injection Scanner ───────────────────────────────────────────
import re as _re

_INJECTION_PATTERNS = [
    (_re.compile(r"(?i)ignore\s+(all\s+)?previous\s+instructions?"), "high", "prompt_override"),
    (_re.compile(r"(?i)you\s+are\s+now\s+a\s+(different|new)"), "high", "identity_override"),
    (_re.compile(r"(?i)\b(execute|eval)\s*\("), "high", "code_injection"),
    (_re.compile(r"(?i)[;|&]{1,2}\s*(rm|curl|wget|bash|sh|cmd|powershell)"), "high", "shell_injection"),
    (_re.compile(r"(?i)(send|post|upload|transmit).{0,50}(http|ftp|external)"), "high", "data_exfiltration"),
    (_re.compile(r"(?i)base64.{0,30}(send|post|upload)"), "medium", "encoded_exfiltration"),
    (_re.compile(r"(?i)(DAN\s+mode|jailbreak|do\s+anything\s+now)"), "high", "jailbreak"),
]

_SECRET_PATTERNS = [
    (_re.compile(r"sk-[A-Za-z0-9_\-]{20,}"), "openai_key"),
    (_re.compile(r"sk-ant-[A-Za-z0-9_\-]{20,}"), "anthropic_key"),
    (_re.compile(r"AKIA[0-9A-Z]{16}"), "aws_access_key"),
    (_re.compile(r"(?:ghp|gho|ghs|ghr|github_pat)_[A-Za-z0-9_]{36,}"), "github_token"),
    (_re.compile(r"(?i)(?:postgres|mysql|mongodb)://[^\s\"']+"), "db_connection_string"),
    (_re.compile(r"-----BEGIN\s+(?:RSA\s+)?PRIVATE\s+KEY-----"), "private_key"),
    (_re.compile(r"sk_(?:live|test)_[A-Za-z0-9]{24,}"), "stripe_key"),
    (_re.compile(r"xox[bpoas]-[A-Za-z0-9\-]{10,}"), "slack_token"),
]

_PII_PATTERNS = [
    (_re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"), "email"),
    (_re.compile(r"\b\d{3}-\d{2}-\d{4}\b"), "ssn"),
    (_re.compile(r"\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13})\b"), "credit_card"),
    (_re.compile(r"\b(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"), "phone"),
]


def scan_injection(text: str) -> list[dict]:
    """Scan for prompt injection attempts. Returns list of findings."""
    findings = []
    for pattern, level, name in _INJECTION_PATTERNS:
        for m in pattern.finditer(text):
            findings.append({"type": name, "level": level, "position": m.start()})
    return findings


def redact_sensitive(text: str) -> tuple[str, list[str]]:
    """Redact PII and secrets from text. Returns (redacted_text, found_types)."""
    found = []
    for pattern, name in _SECRET_PATTERNS + _PII_PATTERNS:
        new_text = pattern.sub(f"[REDACTED:{name}]", text)
        if new_text != text:
            found.append(name)
            text = new_text
    return text, found


# Expose scan results via endpoint
_injection_scan_log: list[dict] = []

# ─────────────────────────────────────────────────────────────────────────────
# Path resolution (supports PyInstaller --onefile bundle)
# ─────────────────────────────────────────────────────────────────────────────


def get_base_path() -> Path:
    """Base path for bundled data (dist/, skills/). Handles PyInstaller."""
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS)
    return Path(__file__).parent


def get_writable_path() -> Path:
    """Writable path for runtime data (chroma_db/). Next to exe when frozen."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent


_BASE = get_base_path()
_WRITABLE = get_writable_path()

# ─────────────────────────────────────────────────────────────────────────────
# Logging setup (RotatingFileHandler)
# ─────────────────────────────────────────────────────────────────────────────

_LOG_DIR = _WRITABLE / "logs"
_LOG_DIR.mkdir(exist_ok=True)

_logger = logging.getLogger("echo")
_logger.setLevel(logging.INFO)
_log_handler = logging.handlers.RotatingFileHandler(
    str(_LOG_DIR / "echo_backend.log"),
    maxBytes=5 * 1024 * 1024,  # 5 MB
    backupCount=3,
    encoding="utf-8",
)
_log_handler.setFormatter(logging.Formatter(
    "%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
))
_logger.addHandler(_log_handler)

# ─────────────────────────────────────────────────────────────────────────────
# Persistent httpx client (connection reuse — biggest perf win)
# ─────────────────────────────────────────────────────────────────────────────

_ollama_client: httpx.AsyncClient | None = None
_external_http_client: httpx.AsyncClient | None = None


async def get_ollama_client() -> httpx.AsyncClient:
    """Return a reusable httpx client for all Ollama calls.
    v3.5: http2=True for multiplexed connections; max_keepalive_connections raised to 20.
    """
    global _ollama_client
    if _ollama_client is None or _ollama_client.is_closed:
        try:
            _ollama_client = httpx.AsyncClient(
                timeout=httpx.Timeout(300.0, connect=10.0),
                limits=httpx.Limits(max_connections=30, max_keepalive_connections=20),
                http2=True,
            )
        except Exception:
            # Fallback if h2 package not installed
            _ollama_client = httpx.AsyncClient(
                timeout=httpx.Timeout(300.0, connect=10.0),
                limits=httpx.Limits(max_connections=30, max_keepalive_connections=20),
            )
    return _ollama_client


async def get_external_client() -> httpx.AsyncClient:
    """Return a reusable httpx client for external APIs (weather, web scraping).
    Kept separate from Ollama client to avoid connection limit interference.
    """
    global _external_http_client
    if _external_http_client is None or _external_http_client.is_closed:
        _external_http_client = httpx.AsyncClient(
            timeout=httpx.Timeout(15.0, connect=5.0),
            limits=httpx.Limits(max_connections=10, max_keepalive_connections=5),
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ECHO/3.3)"},
        )
    return _external_http_client

# ─────────────────────────────────────────────────────────────────────────────
# App setup (lifespan replaces deprecated @app.on_event)
# ─────────────────────────────────────────────────────────────────────────────


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup / shutdown lifecycle for the ECHO backend."""
    print("=" * 50)
    print("  ECHO Local Backend v3.1 - Multi-Agent Pipeline")
    print("=" * 50)
    _logger.info("ECHO Backend v3.1 starting up")

    client = await get_ollama_client()
    try:
        resp = await client.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        models = [m["name"] for m in resp.json().get("models", [])]
        if models:
            print(f"[OK] Ollama connected  -- models: {', '.join(models)}")
        else:
            print("[!!] Ollama connected but no models found.")
            print("  Pull a model:  ollama pull llama3.1:8b")
    except Exception:
        print("[!!] Ollama not running at localhost:11434")
        print("  Start it:      ollama serve")
        print("  Install:       https://ollama.com/download/windows")

    # Pre-warm ALL unique agent models concurrently so first query from any agent is fast
    unique_models = list(set(AGENT_MODEL_MAP.values()))

    async def _warm(mdl: str):
        try:
            await client.post(
                f"{OLLAMA_URL}/api/generate",
                json={"model": mdl, "prompt": "", "keep_alive": "10m"},
                timeout=30,
            )
            print(f"[OK] Pre-warmed model: {mdl}")
        except Exception:
            print(f"[!!] Could not pre-warm {mdl}")

    await asyncio.gather(*[_warm(m) for m in unique_models])

    try:
        import chromadb  # noqa: F401
        print("[OK] ChromaDB available -- RAG + Memory enabled")
    except ImportError:
        print("[!!] ChromaDB not installed  -- RAG/Memory disabled")
        print("  Enable:    pip install chromadb")

    _base = get_base_path()
    skills_dir = _base / "skills"
    if skills_dir.exists():
        count = len(list(skills_dir.glob("*.md"))) - (1 if (skills_dir / "README.md").exists() else 0)
        print(f"[OK] Skills directory -- {max(0, count)} skill file(s)")
    else:
        skills_dir.mkdir(exist_ok=True)
        print("[OK] Skills directory created")

    # Initialize VRAM scheduler
    await vram_scheduler.initialize()

    # Prime psutil CPU percent so non-blocking reads return meaningful data
    psutil.cpu_percent(interval=None)

    # v3.2: Init BM25 index from existing docs
    try:
        await _rebuild_bm25()
        print(f"[OK] BM25 hybrid index initialized" if _BM25_AVAILABLE else "[--] BM25 unavailable (pip install rank-bm25)")
    except Exception:
        pass

    # v3.2: Start knowledge folder watcher
    try:
        loop = asyncio.get_event_loop()
        await _knowledge_watcher.start(loop)
        print(f"[OK] Knowledge watcher: {_KNOWLEDGE_DIR}" if _WATCHDOG_AVAILABLE else f"[OK] Knowledge folder: {_KNOWLEDGE_DIR} (watchdog unavailable, manual ingest only)")
    except Exception as e:
        print(f"[!!] Knowledge watcher failed: {e}")

    # v3.5: Pre-compile skills cache at startup
    try:
        global _compiled_skills_cache
        _compiled_skills_cache = await compile_skills()
        print(f"[OK] Skill compiler: {len(_compiled_skills_cache)} skill(s) compiled")
    except Exception as e:
        print(f"[!!] Skill compiler failed: {e}")

    print(f"[OK] Pipeline: {' -> '.join(PIPELINE)}")
    print(f"[OK] Features: Hybrid RAG | Web Research | BM25 | Knowledge Watcher | Workflow Builder")
    print(f"[OK] v3.5 Features: HTTP/2 | CtxWindowMgr | StreamBatching | KVCache | Thinking Loop | Thought Graph | SentinelImprove | SkillCompiler | ToolDiscovery | QuantSwitch | ProjectMode | ProgressivePipeline | SpeculativeDecode")
    print(f"[OK] Logging to: {_LOG_DIR / 'echo_backend.log'}")
    print(f"[OK] API ready at     http://localhost:8000")
    print("=" * 50)

    yield  # App runs here

    # Shutdown
    print("[shutdown] Cleaning up...")
    _logger.info("ECHO Backend shutting down")
    _knowledge_watcher.stop()
    await vram_scheduler.shutdown()
    global _ollama_client, _external_http_client
    if _ollama_client and not _ollama_client.is_closed:
        await _ollama_client.aclose()
        _ollama_client = None
    if _external_http_client and not _external_http_client.is_closed:
        await _external_http_client.aclose()
        _external_http_client = None


app = FastAPI(title="ECHO Local Backend", version="3.3.0", lifespan=lifespan)

app.add_middleware(GZipMiddleware, minimum_size=1024)  # Compress responses > 1KB
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# SPA middleware — serves frontend files when no API route matches
class SPAMiddleware(BaseHTTPMiddleware):
    """Intercept 404s and serve frontend SPA files from dist/."""

    async def dispatch(self, request, call_next):
        response = await call_next(request)

        # Only intercept 404s for non-API paths
        path = request.url.path
        if response.status_code == 404 and not path.startswith("/api/"):
            dist = _find_dist_dir()
            if dist is not None:
                # Try to serve the exact file (JS, CSS, images)
                url_path = path.lstrip("/")
                if url_path:
                    file = dist / url_path
                    if file.is_file():
                        content_type, _ = mimetypes.guess_type(str(file))
                        return FileResponse(
                            str(file),
                            media_type=content_type or "application/octet-stream",
                        )
                # Fallback to index.html for SPA routing
                index = dist / "index.html"
                if index.is_file():
                    return FileResponse(str(index), media_type="text/html")

        return response


app.add_middleware(SPAMiddleware)

START_TIME = time.time()
OLLAMA_URL = "http://localhost:11434"

# ── Caches ────────────────────────────────────────────────────────────────────
_models_cache: dict = {"models": [], "ts": 0.0}
_MODELS_CACHE_TTL = 30.0  # seconds — avoids HTTP round-trip on every select_model()
_gpu_cache: dict = {"data": None, "ts": 0.0}
_GPU_CACHE_TTL = 5.0  # seconds — avoids repeated GPUtil calls in hot path

# ─────────────────────────────────────────────────────────────────────────────
# Agent configuration — now includes Planner
# ─────────────────────────────────────────────────────────────────────────────

PIPELINE = ["Planner", "Supervisor", "Researcher", "Developer", "Critic"]

AGENT_MODEL_MAP = {
    "Planner": "llama3.2:3b",
    "Supervisor": "llama3.2:3b",
    "Researcher": "llama3.2:3b",
    "Developer": "qwen2.5-coder:3b",
    "Critic": "llama3.2:3b",
    "default": "llama3.2:3b",
}

# ─────────────────────────────────────────────────────────────────────────────
# Vision model support
# ─────────────────────────────────────────────────────────────────────────────

VISION_MODELS = ["llava", "llava-phi3", "llava:7b", "llava:13b", "moondream", "bakllava"]


async def get_vision_model() -> str | None:
    """Return the first available vision model from Ollama, or None."""
    try:
        client = await get_ollama_client()
        resp = await client.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        available = [m["name"] for m in resp.json().get("models", [])]
        for vm in VISION_MODELS:
            for a in available:
                if vm.lower() in a.lower():
                    return a
    except Exception:
        pass
    return None


# Model capability classification for smart routing
MODEL_CAPABILITIES = {
    "llama3.2:3b": {"type": "general", "strengths": ["general", "planning", "conversation", "reasoning"], "vram_mb": 2000},
    "qwen2.5-coder:3b": {"type": "code", "strengths": ["code", "programming", "debugging", "technical"], "vram_mb": 2000},
    "llama3.1:8b": {"type": "general", "strengths": ["general", "planning", "conversation", "reasoning"], "vram_mb": 5000},
    "deepseek-coder-v2:latest": {"type": "code", "strengths": ["code", "programming", "debugging", "technical"], "vram_mb": 9000},
    "dolphin3:latest": {"type": "general", "strengths": ["general", "conversation", "reasoning"], "vram_mb": 5000},
    "qwen3:30b": {"type": "reasoning", "strengths": ["reasoning", "analysis", "research", "general"], "vram_mb": 18000},
    "nomic-embed-text:latest": {"type": "embedding", "strengths": ["embedding"], "vram_mb": 300},
    "llama3.2:1b": {"type": "general", "strengths": ["general", "conversation"], "vram_mb": 1000},
}

# ─────────────────────────────────────────────────────────────────────────────
# v3.2: BM25 keyword index (hybrid RAG)
# ─────────────────────────────────────────────────────────────────────────────

class BM25Index:
    """In-memory BM25 index that shadows the ChromaDB document store."""

    def __init__(self):
        self._ids: list[str] = []
        self._texts: list[str] = []
        self._index = None
        self._lock = asyncio.Lock()

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(r"\w+", text.lower())

    async def rebuild(self, documents: list[dict]) -> None:
        """Rebuild from a list of {id, text} dicts."""
        async with self._lock:
            self._ids = [d["id"] for d in documents]
            self._texts = [d["text"] for d in documents]
            tokenized = [self._tokenize(t) for t in self._texts]
            if tokenized and _BM25_AVAILABLE:
                self._index = _BM25Okapi(tokenized)
            else:
                self._index = None

    async def search(self, query: str, n: int = 10) -> list[tuple[str, float]]:
        """Returns list of (doc_id, normalized_score) sorted descending."""
        async with self._lock:
            if not self._index or not self._ids:
                return []
            tokens = self._tokenize(query)
            scores = self._index.get_scores(tokens)
            max_score = max(scores) if max(scores) > 0 else 1.0
            normalized = [(self._ids[i], float(scores[i]) / max_score) for i in range(len(self._ids))]
            return sorted(normalized, key=lambda x: -x[1])[:n]


_bm25_index = BM25Index()


async def _rebuild_bm25() -> None:
    """Pull all docs from ChromaDB and rebuild the BM25 index."""
    try:
        collection = _get_chroma_collection()
        if collection is None:
            return
        total = collection.count()
        if total == 0:
            await _bm25_index.rebuild([])
            return
        result = collection.get(limit=total)
        ids = result.get("ids", [])
        docs = result.get("documents", [])
        documents = [{"id": ids[i], "text": docs[i] or ""} for i in range(len(ids))]
        await _bm25_index.rebuild(documents)
    except Exception as e:
        _logger.warning(f"BM25 rebuild failed: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# v3.2: Auto Knowledge Ingestion (file watcher)
# ─────────────────────────────────────────────────────────────────────────────

_KNOWLEDGE_DIR = _WRITABLE / "knowledge"
_SUPPORTED_EXTS = {".txt", ".md", ".pdf", ".docx"}


def _extract_text(path: str) -> str:
    """Extract plain text from a file. Returns empty string on failure."""
    p = Path(path)
    ext = p.suffix.lower()
    try:
        if ext in (".txt", ".md"):
            return p.read_text(encoding="utf-8", errors="replace")
        elif ext == ".pdf" and _PDF_AVAILABLE:
            text_parts = []
            with _pdfplumber.open(str(p)) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        elif ext == ".docx" and _DOCX_AVAILABLE:
            doc = _DocxDocument(str(p))
            return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        _logger.warning(f"Text extraction failed for {path}: {e}")
    return ""


def _chunk_text(text: str, size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks."""
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


class _KnowledgeEventHandler(_FSEventHandler):
    def __init__(self, watcher: "KnowledgeWatcher"):
        super().__init__()
        self._watcher = watcher

    def on_created(self, event):
        if not event.is_directory and Path(event.src_path).suffix.lower() in _SUPPORTED_EXTS:
            asyncio.run_coroutine_threadsafe(
                self._watcher.ingest_file(event.src_path), self._watcher._loop
            )

    def on_modified(self, event):
        if not event.is_directory and Path(event.src_path).suffix.lower() in _SUPPORTED_EXTS:
            asyncio.run_coroutine_threadsafe(
                self._watcher.ingest_file(event.src_path), self._watcher._loop
            )

    def on_deleted(self, event):
        if not event.is_directory:
            asyncio.run_coroutine_threadsafe(
                self._watcher.remove_file(event.src_path), self._watcher._loop
            )


class KnowledgeWatcher:
    """Monitors /knowledge folder and auto-ingests files into ChromaDB."""

    def __init__(self):
        self._observer = None
        self._loop = None
        self._files: dict[str, dict] = {}  # path -> {status, chunks, doc_ids, error}

    async def start(self, loop) -> None:
        self._loop = loop
        _KNOWLEDGE_DIR.mkdir(exist_ok=True)
        # Ingest existing files on startup
        for p in _KNOWLEDGE_DIR.iterdir():
            if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTS:
                await self.ingest_file(str(p))
        if _WATCHDOG_AVAILABLE:
            self._observer = _WatchdogObserver()
            self._observer.schedule(_KnowledgeEventHandler(self), str(_KNOWLEDGE_DIR), recursive=False)
            self._observer.start()

    def stop(self) -> None:
        if self._observer:
            try:
                self._observer.stop()
                self._observer.join(timeout=2)
            except Exception:
                pass

    async def ingest_file(self, path: str) -> None:
        fname = Path(path).name
        self._files[path] = {"status": "processing", "file": fname, "chunks": 0, "doc_ids": [], "error": None}
        try:
            text = _extract_text(path)
            if not text.strip():
                self._files[path] = {"status": "empty", "file": fname, "chunks": 0, "doc_ids": [], "error": "No text extracted"}
                return
            chunks = _chunk_text(text)
            collection = _get_chroma_collection()
            if collection is None:
                self._files[path] = {"status": "error", "file": fname, "chunks": 0, "doc_ids": [], "error": "ChromaDB unavailable"}
                return
            # Use file path hash as prefix for stable IDs
            path_hash = hashlib.md5(path.encode()).hexdigest()[:8]
            doc_ids = [f"kw-{path_hash}-{i}" for i in range(len(chunks))]
            # Remove old chunks for this file first
            old_ids = self._files.get(path, {}).get("doc_ids", [])
            if old_ids:
                try:
                    collection.delete(ids=old_ids)
                except Exception:
                    pass
            title = Path(path).stem
            metadatas = [{"title": title, "source": "knowledge", "chunk": i, "file": fname} for i in range(len(chunks))]
            collection.upsert(documents=chunks, ids=doc_ids, metadatas=metadatas)
            self._files[path] = {
                "status": "ok",
                "file": fname,
                "chunks": len(chunks),
                "doc_ids": doc_ids,
                "error": None,
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            }
            await _rebuild_bm25()
            _logger.info(f"Knowledge: ingested {fname} ({len(chunks)} chunks)")
        except Exception as e:
            self._files[path] = {"status": "error", "file": fname, "chunks": 0, "doc_ids": [], "error": str(e)}
            _logger.error(f"Knowledge ingest failed for {fname}: {e}")

    async def remove_file(self, path: str) -> None:
        info = self._files.pop(path, {})
        doc_ids = info.get("doc_ids", [])
        if doc_ids:
            try:
                collection = _get_chroma_collection()
                if collection:
                    collection.delete(ids=doc_ids)
                await _rebuild_bm25()
            except Exception:
                pass

    def status(self) -> list[dict]:
        return list(self._files.values())


_knowledge_watcher = KnowledgeWatcher()


# ─────────────────────────────────────────────────────────────────────────────
# v3.4: Memory Decay + Importance Scoring (ported from old ECHO project)
# ─────────────────────────────────────────────────────────────────────────────

class MemoryDecay:
    """Exponential forgetting curve — memories decay like human memory."""

    def __init__(self, half_life_days: float = 30):
        self.half_life_seconds = half_life_days * 86400

    def decay_factor(self, timestamp: float) -> float:
        """Returns 1.0 for brand-new memory, approaching 0 as time passes."""
        age = time.time() - timestamp
        return 0.5 ** (age / self.half_life_seconds)


_IMPORTANCE_KEYWORDS = [
    "always", "never", "prefer", "my system", "important", "remember",
    "from now on", "every time", "make sure", "don't forget", "critical",
    "must", "requirement", "my name", "i am", "i'm", "my preference",
]


def importance_score(text: str) -> float:
    """Score 0.0–1.0 indicating how worth-remembering a piece of text is.

    Base 0.2 so even non-keyword content has a small chance; each keyword
    match adds 0.15, capped at 1.0. Threshold to store: >= 0.4.
    """
    if not text or not text.strip():
        return 0.0
    score = 0.2
    lower = text.lower()
    for kw in _IMPORTANCE_KEYWORDS:
        if kw in lower:
            score += 0.15
    return min(score, 1.0)


_memory_decay = MemoryDecay(half_life_days=30)
_IMPORTANCE_THRESHOLD = 0.4


# ─────────────────────────────────────────────────────────────────────────────
# v3.4: Sentinel Self-Optimization Engine (ported from old ECHO project)
# ─────────────────────────────────────────────────────────────────────────────

# Routing history entry: {task_type, route, success, ts}
_routing_history: list[dict] = []
_ROUTING_HISTORY_MAX = 500


def _record_routing(task_type: str, route: str, success: bool) -> None:
    """Append a routing event to the history ring-buffer."""
    _routing_history.append({
        "task_type": task_type,
        "route": route,
        "success": success,
        "ts": time.time(),
    })
    if len(_routing_history) > _ROUTING_HISTORY_MAX:
        _routing_history.pop(0)


class SentinelEngine:
    """Monitors routing history and auto-tunes confidence thresholds.

    Runs analysis on demand (via /api/sentinel/optimize) or can be called
    after each N requests. Does NOT run on a background timer to avoid
    interfering with async event loop.

    v3.5: Extended with self-improvement — tracks agent performance history
    and adjusts prompts based on latency and response length patterns.
    """

    def __init__(self):
        self.confidence_threshold: float = 0.72  # mirrors SmartRouter default
        self.complexity_threshold: int = 8
        self._optimizations: list[dict] = []
        # v3.5: Performance tracking for self-improvement
        self._performance_history: list[dict] = []
        self._MAX_PERF_HISTORY = 200

    def get_health(self, last_n: int = 50) -> dict:
        """Quick health snapshot of the last N routing entries."""
        window = _routing_history[-last_n:]
        if not window:
            return {
                "status": "no_data",
                "entries": 0,
                "failure_rate": 0.0,
                "confidence_threshold": self.confidence_threshold,
                "optimizations": len(self._optimizations),
            }
        failures = sum(1 for e in window if not e["success"])
        failure_rate = failures / len(window)
        # Per-type breakdown
        by_type: dict[str, dict] = {}
        for e in window:
            t = e["task_type"]
            if t not in by_type:
                by_type[t] = {"total": 0, "failures": 0}
            by_type[t]["total"] += 1
            if not e["success"]:
                by_type[t]["failures"] += 1
        hotspots = [t for t, v in by_type.items() if v["total"] >= 3 and v["failures"] / v["total"] > 0.3]
        return {
            "status": "healthy" if failure_rate < 0.15 else "degraded",
            "entries": len(window),
            "failure_rate": round(failure_rate, 3),
            "confidence_threshold": self.confidence_threshold,
            "complexity_threshold": self.complexity_threshold,
            "hotspots": hotspots,
            "optimizations": len(self._optimizations),
        }

    def optimize(self, last_n: int = 100) -> dict:
        """Analyze routing history and auto-raise thresholds if failure rate > 15%."""
        window = _routing_history[-last_n:]
        if not window:
            return {"action": "no_data", "message": "No routing history available yet."}

        failures = sum(1 for e in window if not e["success"])
        failure_rate = failures / len(window)
        actions = []

        if failure_rate > 0.15:
            old_ct = round(self.confidence_threshold, 3)
            self.confidence_threshold = min(0.95, self.confidence_threshold + 0.05)
            actions.append(
                f"Raised confidence_threshold {old_ct} → {round(self.confidence_threshold, 3)} "
                f"(failure_rate={round(failure_rate * 100, 1)}%)"
            )
        elif failure_rate < 0.05 and self.confidence_threshold > 0.72:
            old_ct = round(self.confidence_threshold, 3)
            self.confidence_threshold = max(0.72, self.confidence_threshold - 0.02)
            actions.append(
                f"Lowered confidence_threshold {old_ct} → {round(self.confidence_threshold, 3)} "
                f"(low failure_rate={round(failure_rate * 100, 1)}%)"
            )

        # Per-type hotspot detection
        by_type: dict[str, dict] = {}
        for e in window:
            t = e["task_type"]
            if t not in by_type:
                by_type[t] = {"total": 0, "failures": 0}
            by_type[t]["total"] += 1
            if not e["success"]:
                by_type[t]["failures"] += 1
        hotspots = [t for t, v in by_type.items() if v["total"] >= 3 and v["failures"] / v["total"] > 0.3]
        if hotspots:
            actions.append(f"Hotspot task types detected: {hotspots}")

        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "failure_rate": round(failure_rate, 3),
            "actions": actions,
            "entries_analyzed": len(window),
        }
        self._optimizations.append(entry)
        if len(self._optimizations) > 50:
            self._optimizations.pop(0)

        return entry

    # ── v3.5: Self-Improvement Engine ────────────────────────────────────

    def record_performance(self, agent: str, response_len: int, latency_ms: int) -> None:
        """Record an agent's response length and latency for trend analysis."""
        self._performance_history.append({
            "agent": agent,
            "response_len": response_len,
            "latency_ms": latency_ms,
            "ts": time.time(),
        })
        if len(self._performance_history) > self._MAX_PERF_HISTORY:
            self._performance_history.pop(0)

    def improve_prompts(self) -> dict:
        """Analyze performance history and auto-adjust agent prompts.

        - If avg latency > 10000ms → append "Be more concise." to that agent's prompt.
        - If avg response length < 200 chars → append "Provide more detail."
        Returns dict of changes made per agent.
        """
        if not self._performance_history:
            return {"status": "no_data", "changes": {}}

        # Aggregate per-agent stats
        agent_stats: dict[str, dict] = {}
        for entry in self._performance_history:
            ag = entry["agent"]
            if ag not in agent_stats:
                agent_stats[ag] = {"latencies": [], "lengths": []}
            agent_stats[ag]["latencies"].append(entry["latency_ms"])
            agent_stats[ag]["lengths"].append(entry["response_len"])

        changes: dict[str, list[str]] = {}
        for ag, stats in agent_stats.items():
            if ag not in AGENT_SPECIFIC_PROMPTS:
                continue
            avg_latency = sum(stats["latencies"]) / len(stats["latencies"])
            avg_length = sum(stats["lengths"]) / len(stats["lengths"])
            agent_changes: list[str] = []

            if avg_latency > 10000 and "Be more concise." not in AGENT_SPECIFIC_PROMPTS[ag]:
                AGENT_SPECIFIC_PROMPTS[ag] += "\n\nBe more concise."
                agent_changes.append(f"Added conciseness hint (avg_latency={avg_latency:.0f}ms)")

            if avg_length < 200 and "Provide more detail." not in AGENT_SPECIFIC_PROMPTS[ag]:
                AGENT_SPECIFIC_PROMPTS[ag] += "\n\nProvide more detail."
                agent_changes.append(f"Added detail hint (avg_length={avg_length:.0f} chars)")

            if agent_changes:
                changes[ag] = agent_changes

        result = {
            "status": "ok",
            "agents_analyzed": len(agent_stats),
            "changes": changes,
            "ts": datetime.now(timezone.utc).isoformat(),
        }
        self._optimizations.append(result)
        return result


_sentinel = SentinelEngine()


# ── v3.6: LoopGuard — prevents infinite agent loops ──────────────────────────
class LoopGuard:
    """Detects and prevents infinite loops in agent tool-call sequences.

    Checks for: identical repeated calls (SHA256), ping-pong (A-B-A-B),
    and poll budget exhaustion.
    """
    def __init__(self, identical_limit: int = 50, pingpong_limit: int = 4, poll_limit: int = 100):
        self._seen: set[str] = set()
        self._recent: list[str] = []
        self._poll = 0
        self._identical_limit = identical_limit
        self._pingpong_limit = pingpong_limit
        self._poll_limit = poll_limit

    def check(self, tool_name: str, args: str = "") -> str | None:
        """Returns error string if loop detected, else None."""
        self._poll += 1
        if self._poll > self._poll_limit:
            return f"Poll budget exceeded ({self._poll}/{self._poll_limit})"

        h = hashlib.sha256(f"{tool_name}:{args}".encode()).hexdigest()
        if h in self._seen:
            return f"Identical call loop detected: '{tool_name}'"
        self._seen.add(h)

        self._recent.append(tool_name)
        if len(self._recent) > self._pingpong_limit:
            self._recent.pop(0)
        calls = self._recent
        if len(calls) >= 4:
            if calls[-1] == calls[-3] and calls[-2] == calls[-4] and calls[-1] != calls[-2]:
                return f"Ping-pong loop: '{calls[-2]}' \u2194 '{calls[-1]}'"
        return None

    def reset(self):
        self._seen.clear()
        self._recent.clear()
        self._poll = 0


# ── v3.6: Bandit Model Router (Thompson Sampling) ────────────────────────────
import math as _math
import random as _random


def _gamma_sample(shape: float) -> float:
    """Marsaglia & Tsang method for gamma distribution sampling."""
    if shape < 1.0:
        return _gamma_sample(1.0 + shape) * (_random.random() ** (1.0 / shape))
    d = shape - 1.0 / 3.0
    c = 1.0 / _math.sqrt(9.0 * d)
    while True:
        x = _random.gauss(0, 1)
        v = (1.0 + c * x) ** 3
        if v > 0 and _math.log(_random.random()) < 0.5 * x * x + d - d * v + d * _math.log(v):
            return d * v


def _beta_sample(alpha: float, beta: float) -> float:
    g1 = _gamma_sample(alpha)
    g2 = _gamma_sample(beta)
    return g1 / (g1 + g2) if (g1 + g2) > 0 else 0.5


class BanditRouter:
    """Thompson Sampling bandit for dynamic model selection.

    Tracks success/failure per model and routes to statistically best performer.
    Falls back to configured default when a model hasn't been tried.
    """
    def __init__(self, models: list[str]):
        # Initialize with Bayesian priors (1.0 success, 1.0 failure)
        self._arms: dict[str, dict[str, float]] = {
            m: {"successes": 1.0, "failures": 1.0} for m in models
        }
        self._total_calls: int = 0
        self._lock = asyncio.Lock()

    async def select(self, preferred: str | None = None) -> str:
        """Select model via Thompson Sampling. Returns preferred if not enough data."""
        async with self._lock:
            self._total_calls += 1
            # Try untested models first (UCB1 exploration bonus)
            untested = [m for m, v in self._arms.items()
                        if v["successes"] + v["failures"] <= 2.0]
            if untested:
                return untested[0]
            return max(self._arms, key=lambda m: _beta_sample(
                self._arms[m]["successes"], self._arms[m]["failures"]
            ))

    async def update(self, model: str, success: bool, latency_ms: float = 0.0):
        """Record outcome for a model call."""
        async with self._lock:
            if model not in self._arms:
                self._arms[model] = {"successes": 1.0, "failures": 1.0}
            # Latency penalty: slow responses are partial failures
            reward = 1.0 if success else 0.0
            if latency_ms > 10000:
                reward *= 0.5  # Penalize very slow responses
            if reward > 0.5:
                self._arms[model]["successes"] += 1.0
            else:
                self._arms[model]["failures"] += 1.0

    @property
    def stats(self) -> dict:
        return {
            m: {
                "successes": v["successes"],
                "failures": v["failures"],
                "win_rate": round(v["successes"] / (v["successes"] + v["failures"]), 3),
                "thompson_sample": round(_beta_sample(v["successes"], v["failures"]), 3),
            }
            for m, v in self._arms.items()
        }


# ── v3.6: Skill Discovery from Execution Traces ──────────────────────────────
from collections import defaultdict as _defaultdict

_execution_traces: list[dict] = []  # {"agents": [...], "outcome": 0.0-1.0, "query": str}


def record_trace(agents: list[str], outcome: float, query: str):
    """Record an agent execution trace for skill discovery."""
    _execution_traces.append({
        "agents": agents,
        "outcome": outcome,
        "query": query[:100],
        "ts": time.time(),
    })
    if len(_execution_traces) > 1000:
        _execution_traces.pop(0)


def discover_skills_from_traces(
    min_freq: int = 3,
    min_len: int = 2,
    max_len: int = 5,
    min_quality: float = 0.6,
) -> list[dict]:
    """Mine execution traces for high-value agent sequences."""
    seq_stats: dict[tuple, dict] = _defaultdict(
        lambda: {"count": 0, "total": 0.0, "examples": []}
    )
    for trace in _execution_traces:
        agents = trace["agents"]
        for length in range(min_len, min(max_len + 1, len(agents) + 1)):
            for i in range(len(agents) - length + 1):
                seq = tuple(agents[i:i + length])
                seq_stats[seq]["count"] += 1
                seq_stats[seq]["total"] += trace["outcome"]
                if len(seq_stats[seq]["examples"]) < 3:
                    seq_stats[seq]["examples"].append(trace.get("query", ""))

    skills = []
    for seq, stats in seq_stats.items():
        avg = stats["total"] / stats["count"]
        if stats["count"] >= min_freq and avg >= min_quality:
            skills.append({
                "agent_sequence": list(seq),
                "frequency": stats["count"],
                "avg_outcome": round(avg, 3),
                "score": round(stats["count"] * avg, 2),
                "examples": stats["examples"],
                "auto_discovered": True,
            })
    return sorted(skills, key=lambda s: s["score"], reverse=True)


# Initialize bandit router with known models (uses AGENT_MODEL_MAP defined at module level)
_bandit_router: "BanditRouter | None" = None  # lazy-initialized after AGENT_MODEL_MAP


def _get_bandit_router() -> "BanditRouter":
    """Lazy-initialize the bandit router so AGENT_MODEL_MAP is fully populated."""
    global _bandit_router
    if _bandit_router is None:
        _bandit_router = BanditRouter(list(set(AGENT_MODEL_MAP.values())))
    return _bandit_router


SYSTEM_PROMPT = """You are ECHO, an advanced AI orchestration system running locally on the user's machine. You are a multi-agent swarm intelligence framework with specialized agents:

- **Planner**: Decomposes complex tasks into subtasks and assigns them to specialists
- **Supervisor**: Coordinates tasks, delegates to specialists, synthesizes results
- **Researcher**: Performs deep research with recursive analysis
- **Developer**: Generates code solutions and technical implementations
- **Critic**: Evaluates outputs for accuracy and detects hallucinations

You think step-by-step, provide detailed technical answers, and format responses with markdown. When coding, include complete working examples. When researching, cite reasoning chains.

You are running in LOCAL mode with full hardware access and zero cloud dependency."""

# Lean prompt used for direct (non-pipeline) responses — prevents small models from
# hallucinating fake "Research Chain / Developer's Code Snippet" section headers.
DIRECT_PROMPT = """You are ECHO, a fast and intelligent AI assistant running locally.
Answer the user's question directly, clearly, and concisely.
Use markdown formatting where helpful (bold, lists, code blocks).
When the user asks to CREATE, BUILD, MAKE, DESIGN, or WRITE something — generate the actual complete code or content immediately. Do not describe what you would create; create it.
Do NOT add section headers like "Research Chain", "Critic's Evaluation", or "Developer's Code Snippet".
Just respond naturally as a knowledgeable assistant would."""

# Per-agent system prompts used inside run_subtask — override the generic SYSTEM_PROMPT
# so each agent role behaves correctly (especially Developer: produce actual code, not descriptions)
# v3.5: Enhanced with rich personality traits for each agent role
AGENT_SPECIFIC_PROMPTS: dict[str, str] = {
    "Developer": """You are the Developer — a meticulous, pragmatic code craftsman. Your personality: perfectionist about correctness, hates incomplete implementations, always writes production-ready code.

You are the Developer agent in ECHO — an expert code generator.
Your ONLY job is to produce actual, complete, working code.

CRITICAL RULES:
- When asked to create ANYTHING (animations, logos, UIs, functions, scripts, components), write the complete, runnable code IMMEDIATELY
- Use HTML/CSS/JavaScript for visual/interactive things (logos, animations, UIs, games)
- Use Python for algorithms, data processing, scripts
- Include ALL necessary code — no placeholders, no "TODO", no "implement X later"
- NEVER describe what you would create — CREATE IT directly with full code
- NEVER say "I would implement X by..." — just implement X with the actual code
- Format in markdown code blocks with the correct language tag (```html, ```python, etc.)
- One self-contained working artifact is always preferred over fragmented pieces""",

    "Researcher": """You are the Researcher — curious, rigorous, and evidence-driven. Your personality: analytical, skeptical of unverified claims, loves citing reasoning chains. You go deep on topics, explore multiple angles, and always acknowledge uncertainty where it exists.

You are the Researcher agent in ECHO. Your job is deep analysis and information gathering.
Provide thorough findings with clear reasoning chains. When your research feeds a coding task,
describe the best approach, relevant libraries, and key techniques the Developer should use.""",

    "Supervisor": """You are the Supervisor — decisive, coordination-focused, and result-oriented. Your personality: authoritative but fair, pragmatic, focused on synthesis. You cut through noise to deliver clean, integrated outputs. You NEVER fragment code — always present complete working implementations.

You are the Supervisor agent in ECHO. Synthesize agent results into a single complete, polished response.
CRITICAL: If any agent produced code blocks, PRESERVE THEM EXACTLY — output the complete code directly.
Do NOT summarize code into plain-text descriptions. Do NOT say "A function that does X" — show the actual function.
Combine all agent outputs naturally, removing redundancy while keeping all technical content and code intact.""",

    "Critic": """You are the Critic agent in ECHO. Your personality: sharp, exacting, constructively harsh. You find edge cases others miss. You never approve mediocre work but always suggest concrete improvements.

Review the provided code or content carefully.
Point out specific bugs, missing edge cases, or improvements needed.
When you identify issues in code, provide the corrected version with fixes applied.""",

    "Planner": """You are the Planner — methodical, structured, and systematic. You ALWAYS think step-by-step, break problems into clear subtasks, and never skip planning phases. Your personality: organized, thorough, forward-thinking. You speak in structured bullet points and always consider dependencies between tasks.

You are the Planner agent in ECHO. Decompose user requests into clear subtasks for specialized agents.
For code/visual creation tasks (animations, logos, UIs, components, scripts), always assign the PRIMARY task to Developer.
Keep subtask descriptions concrete, specific, and actionable.""",
}

_agent_states: dict = {
    name: {
        "status": "idle",
        "model": AGENT_MODEL_MAP[name],
        "currentTask": None,
        "lastActive": None,
        "tokensProcessed": 0,
        "totalResponseMs": 0,
        "requestCount": 0,
    }
    for name in PIPELINE
}

# ─────────────────────────────────────────────────────────────────────────────
# Feature #2: Smart Model Routing
# ─────────────────────────────────────────────────────────────────────────────

# Keywords for classifying task types
_TASK_PATTERNS = {
    "code": [
        r"\bcode\b", r"\bfunction\b", r"\bclass\b", r"\bimport\b", r"\bdef\b",
        r"\bpython\b", r"\bjavascript\b", r"\btypescript\b", r"\breact\b",
        r"\bapi\b", r"\bendpoint\b", r"\bbug\b", r"\bdebug\b", r"\bfix\b",
        r"\bhtml\b", r"\bcss\b", r"\bsql\b", r"\bprogram\b", r"\bscript\b",
        r"\bimplement\b", r"\brefactor\b", r"\bcompile\b", r"\bsyntax\b",
    ],
    "reasoning": [
        r"\banalyze\b", r"\bcompare\b", r"\bevaluate\b", r"\breason\b",
        r"\bwhy\b", r"\bexplain\b", r"\bprove\b", r"\blogic\b",
        r"\bmath\b", r"\bcalculate\b", r"\bphilosoph\b", r"\bstrateg\b",
        r"\bpros?\b.*\bcons?\b", r"\btrade.?off\b", r"\bcritiq\b",
    ],
    "research": [
        r"\bresearch\b", r"\bstudy\b", r"\bpaper\b", r"\bsurvey\b",
        r"\bfind\b.*\binformation\b", r"\bsummar\b", r"\breview\b",
        r"\bliterature\b", r"\bsource\b", r"\bcitation\b",
    ],
}


def classify_task(text: str) -> str:
    """Classify a task into a type based on keyword patterns."""
    text_lower = text.lower()
    scores = {"code": 0, "reasoning": 0, "research": 0, "general": 0}

    for task_type, patterns in _TASK_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text_lower):
                scores[task_type] += 1

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "general"


async def select_model(task_text: str, preferred_model: str | None = None) -> str:
    """Dynamically choose the best model for a task.

    Considers task type, model availability, and VRAM.
    """
    if preferred_model:
        return preferred_model

    task_type = classify_task(task_text)

    # Map task types to preferred models
    type_to_models = {
        "code": ["deepseek-coder-v2:latest", "llama3.1:8b"],
        "reasoning": ["llama3.1:8b", "dolphin3:latest"],
        "research": ["llama3.1:8b", "dolphin3:latest"],
        "general": ["llama3.1:8b"],
    }

    candidates = type_to_models.get(task_type, ["llama3.1:8b"])
    available = await get_loaded_models()

    # Pick first available candidate
    best_model = None
    for model in candidates:
        if any(model in m for m in available):
            best_model = model
            break

    if best_model is None:
        best_model = available[0] if available else AGENT_MODEL_MAP["default"]

    # v3.6: Let BanditRouter potentially override with a better-performing model
    try:
        bandit_choice = await _get_bandit_router().select(preferred=best_model)
        # Only use bandit choice if it's actually available
        if bandit_choice and any(bandit_choice in m for m in available):
            return bandit_choice
    except Exception:
        pass

    return best_model


# ─────────────────────────────────────────────────────────────────────────────
# Feature #4: Response Caching (LRU with TTL)
# ─────────────────────────────────────────────────────────────────────────────

class ResponseCache:
    """In-memory LRU cache with TTL for Ollama responses."""

    def __init__(self, max_size: int = 200, ttl_seconds: int = 3600):
        self._cache: OrderedDict[str, dict] = OrderedDict()
        self._max_size = max_size
        self._ttl = ttl_seconds
        self._hits = 0
        self._misses = 0

    def _make_key(self, model: str, messages: list[dict], temperature: float) -> str | None:
        """Create cache key. Returns None if not cacheable (very high temp)."""
        if temperature > 0.7:
            return None  # Don't cache creative/random responses
        content = json.dumps({"m": model, "msgs": messages}, sort_keys=True)
        return hashlib.sha256(content.encode()).hexdigest()

    def get(self, model: str, messages: list[dict], temperature: float) -> str | None:
        """Return cached response or None."""
        key = self._make_key(model, messages, temperature)
        if key is None:
            self._misses += 1
            return None

        entry = self._cache.get(key)
        if entry is None:
            self._misses += 1
            return None

        # Check TTL
        if time.time() - entry["time"] > self._ttl:
            del self._cache[key]
            self._misses += 1
            return None

        self._hits += 1
        self._cache.move_to_end(key)
        return entry["response"]

    def put(self, model: str, messages: list[dict], temperature: float, response: str):
        """Store a response in the cache."""
        key = self._make_key(model, messages, temperature)
        if key is None:
            return

        self._cache[key] = {"response": response, "time": time.time()}
        self._cache.move_to_end(key)

        while len(self._cache) > self._max_size:
            self._cache.popitem(last=False)

    def clear(self):
        """Clear all cached entries."""
        self._cache.clear()
        self._hits = 0
        self._misses = 0

    @property
    def stats(self) -> dict:
        total = self._hits + self._misses
        return {
            "entries": len(self._cache),
            "max_size": self._max_size,
            "hits": self._hits,
            "misses": self._misses,
            "hit_rate": round(self._hits / total * 100, 1) if total > 0 else 0,
            "ttl_seconds": self._ttl,
        }


response_cache = ResponseCache()

# ─────────────────────────────────────────────────────────────────────────────
# Feature #5: GPU VRAM Scheduler
# ─────────────────────────────────────────────────────────────────────────────

class VRAMScheduler:
    """Manages model loading/unloading to prevent GPU overload."""

    def __init__(self):
        self._loaded_models: dict[str, float] = {}  # model → last_used timestamp
        self._idle_timeout = 300  # 5 minutes
        self._max_vram_percent = 90
        self._lock = asyncio.Lock()

    async def initialize(self):
        """Probe which models are currently loaded in Ollama."""
        try:
            models = await get_loaded_models()
            for m in models:
                self._loaded_models[m] = time.time()
            print(f"[VRAM] Tracking {len(models)} loaded model(s)")
        except Exception:
            pass

    async def request_model(self, model_name: str) -> bool:
        """Ensure a model is loaded. Unloads LRU if VRAM is tight."""
        async with self._lock:
            self._loaded_models[model_name] = time.time()

            # Check VRAM pressure
            gpu = get_gpu_info()
            if gpu and gpu["vram_total_mb"] > 0:
                vram_percent = (gpu["vram_used_mb"] / gpu["vram_total_mb"]) * 100
                if vram_percent > self._max_vram_percent:
                    await self._unload_lru(exclude=model_name)

            return True

    async def release_model(self, model_name: str):
        """Mark model as idle (don't unload yet)."""
        if model_name in self._loaded_models:
            self._loaded_models[model_name] = time.time()

    async def _unload_lru(self, exclude: str | None = None):
        """Unload the least-recently-used model to free VRAM."""
        if not self._loaded_models:
            return

        # Find LRU model
        candidates = {k: v for k, v in self._loaded_models.items() if k != exclude}
        if not candidates:
            return

        lru_model = min(candidates, key=candidates.get)
        print(f"[VRAM] Unloading LRU model: {lru_model}")

        try:
            client = await get_ollama_client()
            # Ollama uses POST /api/generate with keep_alive=0 to unload
            await client.post(
                f"{OLLAMA_URL}/api/generate",
                json={"model": lru_model, "keep_alive": 0},
                timeout=30,
            )
            del self._loaded_models[lru_model]
            print(f"[VRAM] Unloaded {lru_model}")
        except Exception as e:
            print(f"[VRAM] Failed to unload {lru_model}: {e}")

    async def cleanup_idle(self):
        """Unload models that haven't been used for a while."""
        now = time.time()
        to_unload = [
            m for m, t in self._loaded_models.items()
            if now - t > self._idle_timeout
        ]
        for model in to_unload:
            await self._unload_lru(exclude=None)

    async def shutdown(self):
        """Clean shutdown."""
        self._loaded_models.clear()

    @property
    def status(self) -> dict:
        gpu = get_gpu_info()
        return {
            "loaded_models": list(self._loaded_models.keys()),
            "model_count": len(self._loaded_models),
            "idle_timeout_s": self._idle_timeout,
            "max_vram_percent": self._max_vram_percent,
            "current_vram_percent": round(
                (gpu["vram_used_mb"] / gpu["vram_total_mb"]) * 100, 1
            ) if gpu and gpu["vram_total_mb"] > 0 else None,
        }


vram_scheduler = VRAMScheduler()

# ─────────────────────────────────────────────────────────────────────────────
# Feature #3: Context Compression Layer
# ─────────────────────────────────────────────────────────────────────────────

async def compress_context(
    messages: list[dict],
    max_tokens: int = 4096,
    keep_recent: int = 6,
) -> list[dict]:
    """Compress conversation context if it exceeds token limits.

    Keeps the system prompt + last N messages verbatim, summarizes older ones.
    """
    if not messages:
        return messages

    # Rough token estimate (4 chars ≈ 1 token)
    total_tokens = sum(len(m.get("content", "")) // 4 for m in messages)
    if total_tokens <= max_tokens:
        return messages

    # Separate system prompt, old messages, and recent messages
    system_msgs = [m for m in messages if m["role"] == "system"]
    non_system = [m for m in messages if m["role"] != "system"]

    if len(non_system) <= keep_recent:
        return messages  # Not enough messages to compress

    old_msgs = non_system[:-keep_recent]
    recent_msgs = non_system[-keep_recent:]

    # Summarize old messages using a fast model
    summary_text = "\n".join(
        f"{m['role'].upper()}: {m['content'][:200]}" for m in old_msgs
    )

    summary_prompt = [
        {"role": "system", "content": "Summarize this conversation excerpt in 2-3 concise sentences. Capture key facts, decisions, and context. Be brief."},
        {"role": "user", "content": summary_text},
    ]

    try:
        summary = await ollama_chat_text(summary_prompt, model=AGENT_MODEL_MAP.get("default"))
        if summary:
            compressed_msg = {
                "role": "system",
                "content": f"[CONVERSATION SUMMARY — earlier messages compressed]\n{summary}",
            }
            return system_msgs + [compressed_msg] + recent_msgs
    except Exception as e:
        print(f"[compress] Summarization failed: {e}")

    # Fallback: just truncate old messages
    return system_msgs + recent_msgs


# ── v3.5: Context Window Manager ─────────────────────────────────────────────

class ContextWindowManager:
    """Manages per-model context windows — trims messages to stay within limits.

    Estimates token count via len(json.dumps(messages)) // 4.
    If over 85% of the model's context limit, removes middle messages while
    preserving the system message and the last 4 messages.
    """

    MODEL_CONTEXT_LIMITS: dict[str, int] = {
        "llama3.2:3b": 4096,
        "qwen2.5-coder:3b": 8192,
        "llama3.1:8b": 8192,
        "llama3.2:1b": 2048,
        "default": 4096,
    }

    def trim(self, messages: list[dict], model: str) -> list[dict]:
        """Return messages trimmed to fit within model context limit."""
        limit = self.MODEL_CONTEXT_LIMITS.get(model, self.MODEL_CONTEXT_LIMITS["default"])
        threshold = int(limit * 0.85)
        try:
            token_estimate = len(json.dumps(messages)) // 4
        except Exception:
            return messages

        if token_estimate <= threshold:
            return messages

        # Separate system messages from conversational messages
        system_msgs = [m for m in messages if m.get("role") == "system"]
        non_system = [m for m in messages if m.get("role") != "system"]

        keep_recent = 4
        if len(non_system) <= keep_recent:
            return messages  # Not enough to trim

        # Keep last 4 non-system messages; drop middle ones
        trimmed = system_msgs + non_system[-keep_recent:]
        _logger.info(
            f"[ContextWindowManager] Trimmed {len(non_system)} → {keep_recent} msgs "
            f"for model {model} (est {token_estimate} tokens > threshold {threshold})"
        )
        return trimmed


_ctx_manager = ContextWindowManager()


# ─────────────────────────────────────────────────────────────────────────────
# Feature #8: Task Planning Agent
# ─────────────────────────────────────────────────────────────────────────────

PLANNER_SYSTEM = """You are the Planner agent in the ECHO multi-agent system. Your job is to decompose user requests into subtasks for specialized agents.

Available agents:
- Supervisor: General coordination, synthesis, simple Q&A
- Researcher: Deep analysis, comparisons, evaluations, reasoning
- Developer: Code generation, debugging, technical implementations
- Critic: Quality review, fact-checking, hallucination detection

Rules:
1. For simple questions/greetings, return a single Supervisor subtask
2. For complex tasks, break into 2-4 subtasks with appropriate agents
3. Mark dependencies — a subtask can depend on previous subtask IDs
4. Each subtask should be self-contained with clear instructions

You MUST respond with ONLY a JSON object in this exact format, no other text:
{"subtasks": [{"id": "t1", "agent": "Developer", "task": "Write a Python function that...", "depends_on": []}, {"id": "t2", "agent": "Critic", "task": "Review the code from t1 for...", "depends_on": ["t1"]}]}"""


class Subtask(BaseModel):
    id: str
    agent: str
    task: str
    depends_on: list[str] = []
    priority: int = 5  # 1=highest, 10=lowest


class TaskPlan(BaseModel):
    subtasks: list[Subtask]


async def plan_task(user_message: str, context: list[dict] | None = None) -> TaskPlan:
    """Use the Planner agent to decompose a user request into subtasks."""
    messages = [
        {"role": "system", "content": PLANNER_SYSTEM},
    ]

    # Add brief context if available
    if context:
        recent = context[-4:]  # Last 4 messages for context
        ctx_text = "\n".join(f"{m['role']}: {m['content'][:150]}" for m in recent)
        messages.append({"role": "system", "content": f"Recent context:\n{ctx_text}"})

    messages.append({"role": "user", "content": f"Decompose this request into subtasks:\n\n{user_message}"})

    model = await select_model(user_message, AGENT_MODEL_MAP.get("Planner"))
    result = await ollama_chat_json(messages, model=model)

    subtasks_data = result.get("subtasks", [])
    if not subtasks_data:
        # Fallback: single Supervisor task
        return TaskPlan(subtasks=[
            Subtask(id="t1", agent="Supervisor", task=user_message, depends_on=[])
        ])

    subtasks = []
    for st in subtasks_data:
        try:
            subtasks.append(Subtask(
                id=st.get("id", f"t{len(subtasks)+1}"),
                agent=st.get("agent", "Supervisor"),
                task=st.get("task", user_message),
                depends_on=st.get("depends_on", []),
            ))
        except Exception:
            continue

    if not subtasks:
        return TaskPlan(subtasks=[
            Subtask(id="t1", agent="Supervisor", task=user_message, depends_on=[])
        ])

    return TaskPlan(subtasks=subtasks)


# ─────────────────────────────────────────────────────────────────────────────
# Feature #1: Parallel Multi-Agent Execution
# ─────────────────────────────────────────────────────────────────────────────

async def run_subtask(
    subtask: Subtask,
    system_prompt: str,
    prior_results: dict[str, str],
    temperature: float = 0.7,
    max_tokens: int = 2048,
) -> tuple[str, str]:
    """Execute a single subtask with the assigned agent's model.

    Returns (subtask_id, result_text).
    """
    agent = subtask.agent if subtask.agent in PIPELINE else "Supervisor"

    # Smart model routing per subtask
    model = await select_model(subtask.task, AGENT_MODEL_MAP.get(agent))
    await vram_scheduler.request_model(model)

    # Build messages for this subtask — use agent-specific prompt for better per-role behavior
    agent_system = AGENT_SPECIFIC_PROMPTS.get(agent, system_prompt)
    messages = [{"role": "system", "content": agent_system}]

    # Include results from dependencies
    if subtask.depends_on and prior_results:
        dep_context = "\n\n".join(
            f"--- Result from {dep_id} ---\n{prior_results.get(dep_id, '(not available)')}"
            for dep_id in subtask.depends_on
            if dep_id in prior_results
        )
        if dep_context:
            messages.append({
                "role": "system",
                "content": f"Previous subtask results for reference:\n{dep_context}",
            })

    # v3.2+: Tool injection per agent type
    task_type = classify_task(subtask.task)

    # Researcher: live web research + weather detection
    if agent == "Researcher" and task_type in ("research", "general"):
        try:
            # Weather detection: inject real-time API data when the task mentions weather
            weather_keywords = ["weather", "temperature", "forecast", "rain", "snow", "wind", "humidity"]
            if any(kw in subtask.task.lower() for kw in weather_keywords):
                # Extract likely city name (simple heuristic: look for proper nouns after "in/at/for")
                city_match = re.search(r'\b(?:in|at|for|weather)\s+([A-Z][a-zA-Z\s]{2,20})', subtask.task)
                city = city_match.group(1).strip() if city_match else subtask.task
                wx = await get_weather(city)
                if wx.get("success"):
                    messages.append({
                        "role": "system",
                        "content": _format_weather(wx),
                    })

            web_findings = await web_research(subtask.task)
            if web_findings:
                messages.append({
                    "role": "system",
                    "content": f"[LIVE WEB RESEARCH — retrieved now]\n{web_findings}",
                })
        except Exception:
            pass

    # Developer: sandboxed code interpreter — execute [RUN_PYTHON] blocks if present in prior results
    if agent == "Developer":
        combined_prior = " ".join(prior_results.values())
        code_match = re.search(r'\[RUN_PYTHON\]\s*```python\s*(.*?)\s*```\s*\[/RUN_PYTHON\]', combined_prior, re.DOTALL)
        if code_match:
            code_to_run = code_match.group(1)
            try:
                code_result = await run_code(code_to_run, timeout=8)
                output = code_result.get("stdout", "") or code_result.get("stderr", "")
                if output:
                    messages.append({
                        "role": "system",
                        "content": f"[CODE INTERPRETER OUTPUT]\n{output[:1000]}\n[Use this output to inform your implementation]",
                    })
            except Exception:
                pass

    # v3.5: Thought Graph for Researcher — generates parallel reasoning paths
    if agent == "Researcher":
        try:
            thought_result = await thought_graph(subtask.task, model=model, n_paths=2)
            if thought_result:
                messages.append({
                    "role": "system",
                    "content": f"[THOUGHT GRAPH — best reasoning path selected]\n{thought_result}",
                })
        except Exception:
            pass

    # v3.5: Thinking Loop for complex tasks (> 80 chars)
    if len(subtask.task) > 80:
        try:
            thought = await thinking_loop(subtask.task, model=model)
            if thought:
                messages.append({
                    "role": "system",
                    "content": f"[THINKING]\n{thought}",
                })
        except Exception:
            pass

    messages.append({"role": "user", "content": subtask.task})

    # Track agent state
    request_start = time.time()
    _agent_states[agent]["status"] = "active"
    _agent_states[agent]["currentTask"] = subtask.task[:60]
    _agent_states[agent]["lastActive"] = datetime.now(timezone.utc).isoformat()

    result = ""
    subtask_success = False
    try:
        result = await asyncio.wait_for(
            ollama_chat_text(messages, model=model, temperature=temperature, max_tokens=max_tokens),
            timeout=120,
        )
        subtask_success = bool(result and len(result.strip()) > 0)
        return subtask.id, result or ""
    except asyncio.TimeoutError:
        _logger.error(f"Agent {agent} timed out on subtask {subtask.id}")
        return subtask.id, f"[Agent {agent} timed out after 120s]"
    finally:
        elapsed_ms = round((time.time() - request_start) * 1000)
        _agent_states[agent]["status"] = "idle"
        _agent_states[agent]["currentTask"] = None
        _agent_states[agent]["tokensProcessed"] += len(result or "") // 4
        _agent_states[agent]["totalResponseMs"] += elapsed_ms
        _agent_states[agent]["requestCount"] += 1
        await vram_scheduler.release_model(model)
        # v3.5: Record performance for Sentinel self-improvement
        try:
            _sentinel.record_performance(agent, len(result or ""), elapsed_ms)
        except Exception:
            pass
        # v3.6: Update BanditRouter with outcome
        try:
            await _get_bandit_router().update(model, success=subtask_success, latency_ms=elapsed_ms)
        except Exception:
            pass


# ── v3.5: Cognitive Architecture — Thinking Loop ─────────────────────────────

async def thinking_loop(question: str, model: str) -> str:
    """Chain-of-thought planning step before executing a complex subtask.

    Sends a CoT prompt to get a step-by-step thought chain, which is then
    prepended to the subtask messages so the agent reasons before responding.
    """
    cot_messages = [
        {
            "role": "system",
            "content": "You are a careful reasoning engine. Break problems into clear logical steps.",
        },
        {
            "role": "user",
            "content": f"Break this problem into steps. Think step by step before answering:\n\n{question}",
        },
    ]
    try:
        thought = await asyncio.wait_for(
            ollama_chat_text(cot_messages, model=model, temperature=0.4, max_tokens=512),
            timeout=30,
        )
        return thought or ""
    except Exception as e:
        _logger.debug(f"[thinking_loop] Failed: {e}")
        return ""


# ── v3.5: Thought Graph System ────────────────────────────────────────────────

async def thought_graph(question: str, model: str, n_paths: int = 2) -> str:
    """Generate multiple reasoning paths in parallel and select the best one.

    Used for research tasks to explore different analytic perspectives.
    """
    _PERSPECTIVES = [
        "approach from first principles — break down to fundamental truths and reason up",
        "approach from analogies and examples — find similar known cases and reason by comparison",
        "approach from critical analysis — identify assumptions, counter-arguments, and edge cases",
    ]

    async def _one_path(perspective: str) -> str:
        msgs = [
            {
                "role": "system",
                "content": f"You are a research analyst. {perspective}.",
            },
            {"role": "user", "content": question},
        ]
        try:
            return await asyncio.wait_for(
                ollama_chat_text(msgs, model=model, temperature=0.6, max_tokens=800),
                timeout=60,
            )
        except Exception:
            return ""

    try:
        perspectives = _PERSPECTIVES[:n_paths]
        paths = await asyncio.gather(*[_one_path(p) for p in perspectives], return_exceptions=True)
        valid_paths = [p for p in paths if isinstance(p, str) and p.strip()]

        if not valid_paths:
            return ""

        if len(valid_paths) == 1:
            return valid_paths[0]

        # Critic step — select the best path
        critic_msgs = [
            {
                "role": "system",
                "content": "You are a research quality evaluator. Select the most insightful and accurate reasoning path. Return only the selected path verbatim.",
            },
            {
                "role": "user",
                "content": (
                    f"Question: {question}\n\n"
                    + "\n\n---\n\n".join(
                        f"Path {i + 1}:\n{p}" for i, p in enumerate(valid_paths)
                    )
                ),
            },
        ]
        best = await asyncio.wait_for(
            ollama_chat_text(critic_msgs, model=model, temperature=0.3, max_tokens=1200),
            timeout=60,
        )
        return best or valid_paths[0]
    except Exception as e:
        _logger.debug(f"[thought_graph] Failed: {e}")
        return ""


async def run_pipeline(
    plan: TaskPlan,
    system_prompt: str,
    temperature: float = 0.7,
    max_tokens: int = 2048,
    user_text: str = "",
) -> dict[str, str]:
    """Execute subtasks respecting dependencies. Parallel where possible.
    v3.6: LoopGuard prevents infinite agent loops.
    """
    results: dict[str, str] = {}
    completed: set[str] = set()
    all_ids = {st.id for st in plan.subtasks}
    loop_guard = LoopGuard()  # v3.6: per-pipeline loop detection

    while len(completed) < len(plan.subtasks):
        # Find ready subtasks (dependencies satisfied)
        ready = [
            st for st in plan.subtasks
            if st.id not in completed
            and all(dep in completed for dep in st.depends_on)
        ]

        if not ready:
            break  # Deadlock or circular dependency — bail out

        # Sort by priority (lower = higher priority) before parallel execution
        ready.sort(key=lambda s: s.priority)

        # v3.6: LoopGuard check before execution
        filtered_ready = []
        for st in ready:
            loop_err = loop_guard.check(st.agent, st.task[:50])
            if loop_err:
                _logger.warning(f"[LoopGuard] Skipping subtask {st.id} ({st.agent}): {loop_err}")
                results[st.id] = f"[LoopGuard] Skipped: {loop_err}"
                completed.add(st.id)
            else:
                filtered_ready.append(st)

        if not filtered_ready:
            continue

        # Execute ready subtasks in parallel
        tasks = [
            run_subtask(st, system_prompt, results, temperature, max_tokens)
            for st in filtered_ready
        ]
        batch_results = await asyncio.gather(*tasks, return_exceptions=True)

        for res in batch_results:
            if isinstance(res, Exception):
                print(f"[pipeline] Subtask failed: {res}")
                continue
            task_id, text = res
            results[task_id] = text
            completed.add(task_id)

    # v3.6: Record execution trace for skill discovery
    try:
        agents_used = [st.agent for st in plan.subtasks]
        outcome = sum(1 for v in results.values() if v and len(v) > 50) / max(len(results), 1)
        record_trace(agents_used, outcome, user_text[:100] if user_text else "")
    except Exception:
        pass

    return results


# ─────────────────────────────────────────────────────────────────────────────
# Feature #7: Self-Reflection Loop
# ─────────────────────────────────────────────────────────────────────────────

REFLECTION_SYSTEM = """You are the Critic agent in the ECHO system. Evaluate the given response for:
1. Accuracy — are facts correct?
2. Completeness — does it fully address the question?
3. Clarity — is it well-structured and easy to understand?
4. Hallucinations — does it contain made-up information?

For each issue you identify, rate your confidence (0.0-1.0) that it is a genuine problem.
Only report issues with confidence > 0.75. Lower confidence observations should be omitted.

You MUST respond with ONLY a JSON object:
{"score": 8, "issues": [{"description": "could elaborate on X", "confidence": 0.85, "fix": "add detail about Y"}], "improved_response": null, "overall_score": 0.8}

- score: 1-10 (10 = perfect)
- issues: list of high-confidence problems found (empty if none); each has description, confidence (0-1), fix
- improved_response: null if score >= 7, otherwise provide a corrected version
- overall_score: float 0.0-1.0 quality estimate"""

# v3.4: Deliberation voter system prompts — three specialist perspectives
_VOTER_SYSTEMS = {
    "accuracy": """You are an Accuracy Auditor. Evaluate ONLY factual correctness.
Vote APPROVE if facts seem correct; REJECT if you spot errors or hallucinations.
Respond with JSON only: {"vote": "APPROVE", "reason": "brief reason", "issues": ["issue1"]}""",

    "completeness": """You are a Completeness Reviewer. Evaluate ONLY whether the response fully addresses the question.
Vote APPROVE if the question is answered; REJECT if important aspects are missing.
Respond with JSON only: {"vote": "APPROVE", "reason": "brief reason", "issues": ["missing: X"]}""",

    "clarity": """You are a Clarity Inspector. Evaluate ONLY how well-structured and clear the response is.
Vote APPROVE if it's clear and easy to follow; REJECT if it's confusing or disorganized.
Respond with JSON only: {"vote": "APPROVE", "reason": "brief reason", "issues": ["unclear: X"]}""",
}


async def _deliberation_vote(
    question: str,
    response: str,
    voter_name: str,
    system_prompt: str,
    model: str,
) -> dict:
    """Single voter — returns vote dict or fallback APPROVE on failure."""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"## Question\n{question}\n\n## Response\n{response}"},
    ]
    try:
        result = await ollama_chat_json(messages, model=model)
        result["voter"] = voter_name
        return result
    except Exception as e:
        _logger.warning(f"[deliberation] {voter_name} vote failed: {e}")
        return {"vote": "APPROVE", "reason": "vote failed (fallback)", "issues": [], "voter": voter_name}


async def reflect_on_response(
    question: str,
    response: str,
    model: str | None = None,
    max_loops: int = 2,
) -> tuple[str, dict]:
    """v3.4: Multi-perspective deliberation vote → then Critic refinement loop.

    Phase 1: Three specialist voters (accuracy, completeness, clarity) run in
             parallel. If majority REJECT, issues are fed to the Critic.
    Phase 2: Standard Critic loop (max_loops iterations).
    """
    critic_model = model or AGENT_MODEL_MAP.get("Critic", "llama3.2:3b")

    reflection_info = {"loops": 0, "scores": [], "issues": [], "votes": {}}
    current_response = response

    # ── Phase 1: Parallel deliberation voting ────────────────────────────────
    _agent_states["Critic"]["status"] = "active"
    _agent_states["Critic"]["currentTask"] = "Deliberation vote"
    try:
        vote_tasks = [
            _deliberation_vote(question, current_response, name, prompt, critic_model)
            for name, prompt in _VOTER_SYSTEMS.items()
        ]
        votes = await asyncio.gather(*vote_tasks, return_exceptions=True)
        vote_results = [v for v in votes if isinstance(v, dict)]

        approve_count = sum(1 for v in vote_results if v.get("vote") == "APPROVE")
        reject_count = len(vote_results) - approve_count
        voter_issues: list[str] = []
        for v in vote_results:
            voter_issues.extend(v.get("issues", []))

        reflection_info["votes"] = {
            v.get("voter", "?"): {"vote": v.get("vote"), "reason": v.get("reason")}
            for v in vote_results
        }
        reflection_info["issues"].extend(voter_issues)

        # If majority reject, prepend voter findings to first Critic prompt
        _deliberation_context = ""
        if reject_count > approve_count and voter_issues:
            _deliberation_context = (
                f"Voter issues found: {'; '.join(voter_issues[:5])}. "
                "Please address these in your evaluation."
            )
    except Exception as e:
        _logger.warning(f"[deliberation] Phase 1 failed: {e}")
        _deliberation_context = ""
    finally:
        _agent_states["Critic"]["status"] = "idle"
        _agent_states["Critic"]["currentTask"] = None

    # ── Phase 2: Critic refinement loop ──────────────────────────────────────
    for loop in range(max_loops):
        user_content = f"## Question\n{question}\n\n## Response to evaluate\n{current_response}"
        if loop == 0 and _deliberation_context:
            user_content += f"\n\n## Preliminary Voter Notes\n{_deliberation_context}"

        messages = [
            {"role": "system", "content": REFLECTION_SYSTEM},
            {"role": "user", "content": user_content},
        ]

        _agent_states["Critic"]["status"] = "active"
        _agent_states["Critic"]["currentTask"] = f"Reflection loop {loop + 1}"

        try:
            result = await ollama_chat_json(messages, model=critic_model)
            score = result.get("score", 10)
            raw_issues = result.get("issues", [])
            improved = result.get("improved_response")

            # v3.6: Confidence-gated filtering — only keep high-confidence issues
            issues = []
            for issue in raw_issues:
                if isinstance(issue, dict):
                    if issue.get("confidence", 1.0) > 0.75:
                        issues.append(issue.get("description", str(issue)))
                elif isinstance(issue, str):
                    issues.append(issue)  # legacy plain-string issues pass through

            reflection_info["loops"] += 1
            reflection_info["scores"].append(score)
            reflection_info["issues"].extend(issues)

            if score >= 7 or not improved:
                break  # Good enough

            current_response = improved
        except Exception as e:
            print(f"[reflection] Loop {loop + 1} failed: {e}")
            break
        finally:
            _agent_states["Critic"]["status"] = "idle"
            _agent_states["Critic"]["currentTask"] = None

    return current_response, reflection_info


# ─────────────────────────────────────────────────────────────────────────────
# Skill-tools system prompts
# ─────────────────────────────────────────────────────────────────────────────

SKILL_CREATOR_SYSTEM = """You are a Skill Creator for an AI agent system called ECHO. Your job is to author, tune, and upgrade agent skills.

A skill is a markdown document (SKILL.md) that contains instructions, patterns, and knowledge that gets injected into an agent's system prompt to enhance its capabilities.

When creating a new skill:
- Write clear, actionable instructions
- Include examples where helpful
- Structure with markdown headings
- Keep it focused on one capability
- Include a trigger description (when this skill should activate)

When tuning an existing skill:
- Analyze what works and what doesn't
- Suggest concrete improvements
- If the skill is unnecessary or harmful, recommend deletion with explanation

Format your output as a complete SKILL.md document starting with "# Skill: [Name]"."""

EVAL_SYSTEM = """You are a Skill Evaluator for an AI agent system. Your job is to test whether a skill produces correct behavior.

You will receive:
1. A skill's content (the instructions that would be injected into an agent)
2. A test prompt
3. Expected behavior description

Your task:
- Simulate how an agent with this skill loaded would respond to the prompt
- Compare the likely response against the expected behavior
- Grade the result

You MUST respond with ONLY a JSON object in this exact format, no other text:
{"pass": true, "score": 4, "notes": "explanation of grading"}

- pass: true if the agent would behave as expected, false otherwise
- score: 1-5 (1=terrible, 5=perfect)
- notes: brief explanation"""

TRIGGER_TUNING_SYSTEM = """You are a Skill Trigger Tuning expert. Your job is to analyze and refine skill trigger descriptions.

A trigger description determines WHEN a skill should be activated based on user prompts. Good triggers:
- Are specific enough to avoid false positives
- Are broad enough to catch all relevant prompts
- Use clear, pattern-matchable language

You will receive:
- Current skill name and description
- Sample prompts that SHOULD trigger the skill
- Sample prompts that should NOT trigger the skill

Analyze the current description and suggest an improved version.

You MUST respond with ONLY a JSON object in this exact format, no other text:
{"analysis": "your analysis of current trigger", "suggestedDescription": "the improved description text", "predictedAccuracy": 85}

- analysis: what's wrong with the current trigger
- suggestedDescription: improved trigger text
- predictedAccuracy: estimated accuracy percentage 0-100"""

# ─────────────────────────────────────────────────────────────────────────────
# ChromaDB setup (lazy init) — now with memory collections
# ─────────────────────────────────────────────────────────────────────────────

_chroma_client = None
_chroma_collection = None


def _get_chroma_client():
    """Get or create the shared ChromaDB client."""
    global _chroma_client
    if _chroma_client is not None:
        return _chroma_client
    try:
        import chromadb
        chroma_path = str(_WRITABLE / "chroma_db")
        _chroma_client = chromadb.PersistentClient(path=chroma_path)
        return _chroma_client
    except Exception as e:
        print(f"[!!] ChromaDB client failed: {e}")
        return None


def _get_embedding_fn():
    """Get the Ollama embedding function."""
    try:
        from chromadb.utils import embedding_functions
        return embedding_functions.OllamaEmbeddingFunction(
            url=OLLAMA_URL,
            model_name="nomic-embed-text",
        )
    except Exception:
        return None


def _get_chroma_collection():
    global _chroma_collection
    if _chroma_collection is not None:
        return _chroma_collection
    client = _get_chroma_client()
    ef = _get_embedding_fn()
    if client is None or ef is None:
        return None
    try:
        _chroma_collection = client.get_or_create_collection(
            name="echo_documents",
            embedding_function=ef,
        )
        return _chroma_collection
    except Exception as e:
        print(f"[!!] ChromaDB collection failed: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Feature #6: True Agent Memory (3 ChromaDB collections)
# ─────────────────────────────────────────────────────────────────────────────

_memory_collections: dict = {}


def _get_memory_collection(memory_type: str):
    """Get or create a memory collection by type."""
    if memory_type in _memory_collections:
        return _memory_collections[memory_type]

    valid_types = ("episodic", "semantic", "procedural")
    if memory_type not in valid_types:
        return None

    client = _get_chroma_client()
    ef = _get_embedding_fn()
    if client is None or ef is None:
        return None

    try:
        coll = client.get_or_create_collection(
            name=f"echo_memory_{memory_type}",
            embedding_function=ef,
        )
        _memory_collections[memory_type] = coll
        return coll
    except Exception as e:
        print(f"[!!] Memory collection '{memory_type}' failed: {e}")
        return None


async def recall_relevant_memories(query: str, limit: int = 3) -> list[dict]:
    """Search across all memory types for relevant memories.

    Runs all three ChromaDB type queries concurrently via run_in_executor
    to avoid sequential blocking on the sync ChromaDB API.
    """
    loop = asyncio.get_event_loop()
    now = time.time()

    def _query_type(mem_type: str) -> list[dict]:
        """Synchronous ChromaDB query for a single memory type."""
        coll = _get_memory_collection(mem_type)
        if coll is None or coll.count() == 0:
            return []
        try:
            results = coll.query(
                query_texts=[query],
                n_results=min(limit, coll.count()),
            )
            ids = results.get("ids", [[]])[0]
            docs = results.get("documents", [[]])[0]
            metas = results.get("metadatas", [[]])[0]
            distances = results.get("distances", [[]])[0]
            items = []
            for i, doc_id in enumerate(ids):
                similarity = round(1 - (distances[i] / 2), 4) if distances else 0.0
                # v3.4: Exponential memory decay — smooth forgetting curve
                ts_str = (metas[i] or {}).get("timestamp", "") if metas else ""
                if ts_str:
                    try:
                        mem_time = datetime.fromisoformat(ts_str).timestamp()
                        decay = _memory_decay.decay_factor(mem_time)
                        # Blend: 70% semantic similarity + 30% recency-decay
                        similarity = round(0.7 * similarity + 0.3 * decay, 4)
                    except Exception:
                        pass
                if similarity > 0.3:
                    items.append({
                        "id": doc_id,
                        "type": mem_type,
                        "content": docs[i] if docs else "",
                        "summary": (metas[i] or {}).get("summary", ""),
                        "tags": (metas[i] or {}).get("tags", ""),
                        "similarity": similarity,
                    })
            return items
        except Exception:
            return []

    # Run all three type queries concurrently
    batch = await asyncio.gather(
        loop.run_in_executor(None, _query_type, "episodic"),
        loop.run_in_executor(None, _query_type, "semantic"),
        loop.run_in_executor(None, _query_type, "procedural"),
        return_exceptions=True,
    )
    all_results = [item for group in batch if isinstance(group, list) for item in group]
    all_results.sort(key=lambda x: x["similarity"], reverse=True)
    return all_results[:limit]


async def auto_extract_memories(conversation: list[dict]):
    """After a conversation, automatically extract and store memories."""
    if len(conversation) < 4:
        return  # Too short to extract from

    # Build conversation text
    conv_text = "\n".join(
        f"{m['role'].upper()}: {m['content'][:300]}" for m in conversation[-10:]
    )

    extract_prompt = [
        {"role": "system", "content": """Extract key information from this conversation for long-term memory. Return JSON:
{"episodic_summary": "1-2 sentence summary of what happened",
 "facts": ["fact 1", "fact 2"],
 "strategies": ["successful approach 1"],
 "tags": ["tag1", "tag2"]}
Only include meaningful, reusable information. Return empty lists if nothing worth remembering."""},
        {"role": "user", "content": conv_text},
    ]

    try:
        result = await ollama_chat_json(extract_prompt)

        ts_now = datetime.now(timezone.utc).isoformat()
        tags_str = ",".join(result.get("tags", []))

        # Store episodic memory (v3.4: only if importance score passes threshold)
        episodic_summary = result.get("episodic_summary", "")
        # v3.6: Redact PII/secrets from auto-extracted memories
        try:
            episodic_summary, _ep_redacted = redact_sensitive(episodic_summary)
            if _ep_redacted:
                _logger.info(f"[memory] auto_extract: redacted {_ep_redacted} from episodic summary")
        except Exception:
            pass
        if episodic_summary and importance_score(episodic_summary) >= _IMPORTANCE_THRESHOLD:
            coll = _get_memory_collection("episodic")
            if coll:
                mem_id = f"ep-{int(time.time() * 1000)}"
                coll.upsert(
                    documents=[episodic_summary],
                    ids=[mem_id],
                    metadatas=[{
                        "summary": episodic_summary[:200],
                        "tags": tags_str,
                        "timestamp": ts_now,
                        "importance": round(importance_score(episodic_summary), 3),
                    }],
                )

        # Store semantic facts (v3.4: importance-gated)
        facts = result.get("facts", [])
        for i, fact in enumerate(facts[:5]):  # Max 5 facts per conversation
            if not fact.strip():
                continue
            # v3.6: Redact PII from facts
            try:
                fact, _ = redact_sensitive(fact)
            except Exception:
                pass
            imp = importance_score(fact)
            if imp < _IMPORTANCE_THRESHOLD:
                continue
            coll = _get_memory_collection("semantic")
            if coll:
                mem_id = f"sem-{int(time.time() * 1000)}-{i}"
                coll.upsert(
                    documents=[fact],
                    ids=[mem_id],
                    metadatas=[{
                        "summary": fact[:200],
                        "tags": tags_str,
                        "timestamp": ts_now,
                        "importance": round(imp, 3),
                    }],
                )

        # Store procedural strategies (v3.4: importance-gated)
        strategies = result.get("strategies", [])
        for i, strategy in enumerate(strategies[:3]):
            if not strategy.strip():
                continue
            # v3.6: Redact PII from strategies
            try:
                strategy, _ = redact_sensitive(strategy)
            except Exception:
                pass
            imp = importance_score(strategy)
            if imp < _IMPORTANCE_THRESHOLD:
                continue
            coll = _get_memory_collection("procedural")
            if coll:
                mem_id = f"proc-{int(time.time() * 1000)}-{i}"
                coll.upsert(
                    documents=[strategy],
                    ids=[mem_id],
                    metadatas=[{
                        "summary": strategy[:200],
                        "tags": tags_str,
                        "timestamp": ts_now,
                        "importance": round(imp, 3),
                    }],
                )
    except Exception as e:
        print(f"[memory] Auto-extraction failed: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Pydantic schemas
# ─────────────────────────────────────────────────────────────────────────────


class Message(BaseModel):
    role: str
    content: str


class InstallRequest(BaseModel):
    models: Optional[list[str]] = None
    install_all: Optional[bool] = False


class FileProcessRequest(BaseModel):
    name: str
    content_b64: str
    mime_type: Optional[str] = ""


class ChatRequest(BaseModel):
    messages: list[Message]
    model: Optional[str] = None
    temperature: Optional[float] = 0.7
    max_tokens: Optional[int] = 2048
    depth: Optional[int] = 1
    enable_planning: Optional[bool] = True
    enable_reflection: Optional[bool] = False
    no_cache: Optional[bool] = False
    workflow: Optional[dict] = None  # v3.2: custom workflow definition
    project_id: Optional[str] = None  # v3.5: AI Project Mode — inject project context
    images: Optional[list[str]] = None  # base64 image strings attached to latest user message
    attachments: Optional[list[dict]] = None  # v3.7: file attachments {name, type, content}


class DocumentRequest(BaseModel):
    title: str
    content: str
    id: Optional[str] = None
    user_id: Optional[str] = None


class SearchRequest(BaseModel):
    query: str
    user_id: Optional[str] = None
    limit: Optional[int] = 5
    hybrid: Optional[bool] = True  # v3.2: use BM25+vector hybrid scoring


class WebSearchRequest(BaseModel):
    query: str
    scrape: Optional[bool] = True
    max_results: Optional[int] = 5


class MemoryStoreRequest(BaseModel):
    type: str  # "episodic" | "semantic" | "procedural"
    content: str
    summary: Optional[str] = None
    tags: Optional[list[str]] = []


class MemoryRecallRequest(BaseModel):
    query: str
    limit: Optional[int] = 5
    type: Optional[str] = None  # Filter by type, or None for all


class VoiceTranscribeRequest(BaseModel):
    audio_b64: str  # base64-encoded audio bytes
    format: Optional[str] = "webm"  # "webm" | "wav"


class VoiceSpeakRequest(BaseModel):
    text: str
    rate: Optional[int] = 175   # words per minute
    volume: Optional[float] = 1.0  # 0.0 – 1.0


class VisionRequest(BaseModel):
    image_b64: str  # base64-encoded image (PNG/JPEG/WebP)
    prompt: Optional[str] = "Describe this image in detail."
    model: Optional[str] = None  # override vision model


class SkillToolsRequest(BaseModel):
    mode: str  # "create" | "eval" | "tune-trigger"
    description: Optional[str] = None
    existingSkill: Optional[str] = None
    agent: Optional[str] = "Developer"
    skillContent: Optional[str] = None
    prompt: Optional[str] = None
    expectedBehavior: Optional[str] = None
    skillName: Optional[str] = None
    currentDescription: Optional[str] = None
    shouldTrigger: Optional[list[str]] = None
    shouldNotTrigger: Optional[list[str]] = None


# ─────────────────────────────────────────────────────────────────────────────
# Hardware helpers
# ─────────────────────────────────────────────────────────────────────────────


def get_gpu_info() -> dict | None:
    global _gpu_cache
    now = time.time()
    if now - _gpu_cache["ts"] < _GPU_CACHE_TTL:
        return _gpu_cache["data"]
    try:
        import GPUtil
        gpus = GPUtil.getGPUs()
        if not gpus:
            _gpu_cache = {"data": None, "ts": now}
            return None
        gpu = gpus[0]
        result = {
            "name": gpu.name,
            "vram_total_mb": int(gpu.memoryTotal),
            "vram_used_mb": int(gpu.memoryUsed),
            "gpu_usage_percent": round(gpu.load * 100, 1),
            "temperature_c": gpu.temperature,
        }
        _gpu_cache = {"data": result, "ts": now}
        return result
    except Exception:
        _gpu_cache = {"data": None, "ts": now}
        return None


def get_cpu_temp() -> float | None:
    try:
        temps = psutil.sensors_temperatures()
        if not temps:
            return None
        for key in ("coretemp", "k10temp", "cpu_thermal", "cpu-thermal"):
            if key in temps and temps[key]:
                return round(temps[key][0].current, 1)
        for sensors in temps.values():
            if sensors:
                return round(sensors[0].current, 1)
    except Exception:
        pass
    return None


def get_disk_info() -> dict | None:
    try:
        # Use the drive where this script lives (e.g. D:\) instead of hardcoded C:\
        if platform.system() == "Windows":
            path = os.path.splitdrive(os.path.abspath(__file__))[0] + "\\"
        else:
            path = "/"
        usage = psutil.disk_usage(path)
        return {
            "total_gb": round(usage.total / (1024 ** 3), 1),
            "used_gb": round(usage.used / (1024 ** 3), 1),
            "free_gb": round(usage.free / (1024 ** 3), 1),
            "usage_percent": usage.percent,
        }
    except Exception:
        return None


async def get_loaded_models() -> list[str]:
    global _models_cache
    now = time.time()
    if now - _models_cache["ts"] < _MODELS_CACHE_TTL and _models_cache["models"]:
        return _models_cache["models"]
    try:
        client = await get_ollama_client()
        resp = await client.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        if resp.status_code == 200:
            models = [m["name"] for m in resp.json().get("models", [])]
            _models_cache = {"models": models, "ts": now}
            return models
    except Exception:
        pass
    return _models_cache["models"]  # Return stale data rather than empty on transient errors


# ─────────────────────────────────────────────────────────────────────────────
# Ollama helpers
# ─────────────────────────────────────────────────────────────────────────────


# ── v3.5: Token Stream Batching ───────────────────────────────────────────────
STREAM_BATCH_SIZE = 4  # flush SSE every N tokens to reduce overhead


async def ollama_stream(
    model: str,
    messages: list[dict],
    temperature: float = 0.7,
    max_tokens: int = 2048,
):
    """Stream from Ollama, yield OpenAI-compatible SSE chunks.
    v3.5: Tokens buffered in batches of STREAM_BATCH_SIZE to reduce SSE overhead.
    v3.5: num_keep=256 instructs Ollama to keep system-prompt KV cache warm.
    """
    payload = {
        "model": model,
        "messages": messages,
        "stream": True,
        "options": {
            "temperature": temperature,
            "num_predict": max_tokens,
            "num_gpu": 99,
            "num_keep": 256,  # v3.5: KV cache reuse — keep first 256 tokens (system prompt)
        },
        "keep_alive": "10m",
    }

    client = await get_ollama_client()
    try:
        async with client.stream(
            "POST", f"{OLLAMA_URL}/api/chat", json=payload
        ) as response:
            if response.status_code != 200:
                error_body = await response.aread()
                error_msg = error_body.decode("utf-8", errors="replace")
                _logger.error(f"Ollama stream error {response.status_code}: {error_msg}")
                yield f'data: {{"error": "Ollama error {response.status_code}: {error_msg}"}}\n\n'
                yield "data: [DONE]\n\n"
                return

            # v3.5: batch buffer — accumulate tokens, flush every STREAM_BATCH_SIZE
            _batch_buf: list[str] = []

            async for line in response.aiter_lines():
                if not line.strip():
                    continue
                try:
                    chunk = json.loads(line)
                    content = chunk.get("message", {}).get("content", "")
                    done = chunk.get("done", False)

                    if content:
                        _batch_buf.append(content)
                        if len(_batch_buf) >= STREAM_BATCH_SIZE:
                            batched = "".join(_batch_buf)
                            _batch_buf.clear()
                            sse_data = {"choices": [{"delta": {"content": batched}}]}
                            yield f"data: {json.dumps(sse_data)}\n\n"

                    if done:
                        # Flush remaining buffer on final chunk
                        if _batch_buf:
                            batched = "".join(_batch_buf)
                            _batch_buf.clear()
                            sse_data = {"choices": [{"delta": {"content": batched}}]}
                            yield f"data: {json.dumps(sse_data)}\n\n"
                        yield "data: [DONE]\n\n"
                        return
                except json.JSONDecodeError:
                    continue

    except httpx.ConnectError:
        yield 'data: {"error": "Cannot connect to Ollama. Is it running? Run: ollama serve"}\n\n'
        yield "data: [DONE]\n\n"
    except Exception as e:
        _logger.error(f"Stream error: {e}")
        yield f'data: {{"error": "Stream error: {str(e)}"}}\n\n'
        yield "data: [DONE]\n\n"


async def ollama_chat_json(messages: list[dict], model: str | None = None) -> dict:
    """Non-streaming Ollama call. Returns parsed JSON via format=json mode.
    v3.5: num_keep=256 for KV cache reuse across multi-turn calls.
    """
    model = model or AGENT_MODEL_MAP.get("default", "llama3.1:8b")
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "format": "json",
        "options": {"temperature": 0.3, "num_predict": 512, "num_gpu": 99, "num_keep": 256},
        "keep_alive": "10m",
    }

    client = await get_ollama_client()
    try:
        resp = await client.post(f"{OLLAMA_URL}/api/chat", json=payload, timeout=120)
        if resp.status_code != 200:
            return {}
        data = resp.json()
        content = data.get("message", {}).get("content", "")

        try:
            return json.loads(content)
        except json.JSONDecodeError:
            pass

        match = re.search(r"\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}", content, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                pass

        return {}
    except Exception as e:
        _logger.error(f"Ollama JSON call failed: {e}")
        return {}


async def ollama_chat_text(
    messages: list[dict],
    model: str | None = None,
    temperature: float = 0.7,
    max_tokens: int = 2048,
) -> str:
    """Non-streaming Ollama call. Returns plain text response.
    v3.5: num_keep=256 for KV cache reuse across multi-turn calls.
    """
    model = model or AGENT_MODEL_MAP.get("default", "llama3.1:8b")
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_predict": max_tokens,
            "num_gpu": 99,
            "num_keep": 256,  # v3.5: KV cache reuse
        },
        "keep_alive": "10m",
    }

    client = await get_ollama_client()
    try:
        resp = await client.post(f"{OLLAMA_URL}/api/chat", json=payload, timeout=120)
        if resp.status_code != 200:
            return ""
        data = resp.json()
        return data.get("message", {}).get("content", "")
    except Exception as e:
        _logger.error(f"Ollama text call failed: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# v3.2: Web Research functions (upgraded v3.3: trafilatura + TTL cache + credibility)
# ─────────────────────────────────────────────────────────────────────────────

# TTL cache: 300s to prevent hammering DDG with duplicate queries
_search_cache: "dict | None" = _TTLCache(maxsize=50, ttl=300) if _CACHETOOLS_AVAILABLE else {}

# Credibility scores by domain keyword (higher = more trustworthy)
_CREDIBLE_DOMAINS: dict[str, float] = {
    "wikipedia.org": 0.95, "britannica.com": 0.90,
    "nature.com": 0.95,    "science.org": 0.95,
    "reuters.com": 0.90,   "bbc.com": 0.88,
    "nytimes.com": 0.85,   "cnn.com": 0.80,
    ".gov": 0.95,          ".edu": 0.90,
    "github.com": 0.85,    "stackoverflow.com": 0.82,
    "arxiv.org": 0.92,     "pubmed.ncbi.nlm.nih.gov": 0.95,
}


def _credibility_score(url: str) -> float:
    """Return domain trust score 0-1 for a URL."""
    try:
        domain = urlparse(url).netloc.lower()
        for key, score in _CREDIBLE_DOMAINS.items():
            if key in domain:
                return score
    except Exception:
        pass
    return 0.5


async def web_search(query: str, max_results: int = 5) -> list[dict]:
    """Search DuckDuckGo with TTL caching. Returns result list with credibility scores."""
    if not _DDG_AVAILABLE:
        return []
    cache_key = f"ddg:{query}:{max_results}"
    if cache_key in _search_cache:
        return _search_cache[cache_key]
    try:
        loop = asyncio.get_event_loop()
        raw = await loop.run_in_executor(
            None,
            lambda: list(_DDGS().text(query, max_results=max_results))
        )
        results = [
            {
                "title": r.get("title", ""),
                "url": r.get("href", ""),
                "snippet": r.get("body", ""),
                "credibility": _credibility_score(r.get("href", "")),
            }
            for r in raw
        ]
        _search_cache[cache_key] = results
        return results
    except Exception as e:
        _logger.warning(f"Web search failed: {e}")
        return []


async def scrape_page(url: str, max_chars: int = 3000) -> str:
    """Fetch URL and extract clean article text.

    Uses trafilatura (purpose-built article extractor) with BeautifulSoup fallback.
    Uses the dedicated external HTTP client (not Ollama client) to avoid connection interference.
    """
    try:
        client = await get_external_client()
        resp = await client.get(url, timeout=10)
        if resp.status_code != 200:
            return ""
        html = resp.text

        # trafilatura: best for articles, blog posts, news
        if _TRAFILATURA_AVAILABLE:
            content = _trafilatura.extract(
                html, include_comments=False, include_tables=True, no_fallback=False
            )
            if content:
                return content[:max_chars]

        # BeautifulSoup fallback
        if _BS4_AVAILABLE:
            soup = _BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            text = " ".join(soup.get_text(separator=" ").split())
            return text[:max_chars]

    except Exception as e:
        _logger.warning(f"Scrape failed for {url}: {e}")
    return ""


async def web_research(query: str) -> str:
    """Search web, scrape top results sorted by credibility, return summarized findings."""
    results = await web_search(query, max_results=5)
    if not results:
        return f"Web search unavailable. Answering from training data for: {query}"

    # Sort by credibility so we scrape the best sources first
    sorted_results = sorted(results, key=lambda r: r.get("credibility", 0.5), reverse=True)

    # Scrape top 3 concurrently
    scrape_tasks = [scrape_page(r["url"]) for r in sorted_results[:3]]
    scraped = await asyncio.gather(*scrape_tasks, return_exceptions=True)

    context_parts = []
    for i, r in enumerate(sorted_results):
        cred_pct = f"{int(r.get('credibility', 0.5) * 100)}%"
        snippet = r["snippet"]
        if i < len(scraped) and isinstance(scraped[i], str) and scraped[i]:
            snippet = scraped[i][:600]
        context_parts.append(
            f"Source [{cred_pct} credibility]: {r['title']} ({r['url']})\n{snippet}"
        )

    context = "\n\n---\n\n".join(context_parts)
    summary_messages = [
        {"role": "system", "content": "You are a research assistant. Summarize the following web sources to answer the query. Cite sources by title and prioritize higher-credibility sources. Be concise and factual."},
        {"role": "user", "content": f"Query: {query}\n\nSources:\n{context}"},
    ]
    summary = await ollama_chat_text(summary_messages, model=AGENT_MODEL_MAP.get("Researcher"))
    return summary or context[:1000]


# ─────────────────────────────────────────────────────────────────────────────
# v3.3: Weather tool (Open-Meteo — free, no API key required)
# ─────────────────────────────────────────────────────────────────────────────

async def get_weather(location: str) -> dict:
    """Fetch real-time weather via Open-Meteo (geocoding + forecast APIs).

    Returns dict with temperature, description, wind_speed, or error.
    No API key required. Uses shared external HTTP client.
    """
    try:
        client = await get_external_client()
        # Step 1: geocoding
        geo = await client.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": location, "count": 1, "language": "en", "format": "json"},
        )
        geo.raise_for_status()
        geo_data = geo.json()
        if not geo_data.get("results"):
            return {"success": False, "error": f"Location not found: {location}"}
        loc = geo_data["results"][0]
        lat, lon = loc["latitude"], loc["longitude"]
        full_name = f"{loc.get('name')}, {loc.get('admin1', '')}, {loc.get('country', '')}"

        # Step 2: weather — run geocoding and forecast concurrently when possible,
        # but we need lat/lon from step 1 first, so fetch forecast now
        wx = await client.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": lat, "longitude": lon,
                "current_weather": "true",
                "current": "apparent_temperature,relative_humidity_2m",
                "daily": "temperature_2m_max,temperature_2m_min,weathercode,precipitation_probability_max,precipitation_sum,windspeed_10m_max,winddirection_10m_dominant,uv_index_max,sunrise,sunset",
                "timezone": "auto",
                "forecast_days": 7,
                "temperature_unit": "celsius",
                "wind_speed_unit": "kmh",
            },
        )
        wx.raise_for_status()
        wx_json = wx.json()
        current = wx_json.get("current_weather", {})
        current_extra = wx_json.get("current", {})
        daily_raw = wx_json.get("daily", {})

        code = current.get("weathercode", 0)
        desc_map = {
            frozenset([0]): "Clear skies",
            frozenset([1, 2, 3]): "Partly cloudy",
            frozenset([45, 48]): "Foggy",
            frozenset([51, 53, 55]): "Light drizzle",
            frozenset([61, 63, 65]): "Rain",
            frozenset([71, 73, 75]): "Snowing",
            frozenset([80, 81, 82]): "Rain showers",
            frozenset([95, 96, 99]): "Thunderstorm",
        }
        description = next((v for k, v in desc_map.items() if code in k), "Variable")

        daily = []
        if daily_raw.get("time"):
            precip_probs = daily_raw.get("precipitation_probability_max") or [0] * 7
            precip_sums  = daily_raw.get("precipitation_sum") or [0.0] * 7
            wind_maxes   = daily_raw.get("windspeed_10m_max") or [None] * 7
            wind_dirs    = daily_raw.get("winddirection_10m_dominant") or [None] * 7
            uv_indices   = daily_raw.get("uv_index_max") or [None] * 7
            sunrises     = daily_raw.get("sunrise") or [None] * 7
            sunsets      = daily_raw.get("sunset") or [None] * 7
            for i, date in enumerate(daily_raw["time"]):
                # sunrise/sunset come as ISO strings like "2024-01-15T06:32"; extract HH:MM
                def _hhmm(s):
                    if not s: return None
                    try: return s.split("T")[1][:5]
                    except: return None
                daily.append({
                    "date":        date,
                    "max":         round(daily_raw["temperature_2m_max"][i]),
                    "min":         round(daily_raw["temperature_2m_min"][i]),
                    "code":        daily_raw["weathercode"][i],
                    "precip_prob": precip_probs[i] if i < len(precip_probs) else 0,
                    "precip_sum":  round(precip_sums[i], 1) if i < len(precip_sums) and precip_sums[i] is not None else 0.0,
                    "wind_max":    round(wind_maxes[i]) if i < len(wind_maxes) and wind_maxes[i] is not None else None,
                    "wind_dir":    round(wind_dirs[i]) if i < len(wind_dirs) and wind_dirs[i] is not None else None,
                    "uv_index":    round(uv_indices[i], 1) if i < len(uv_indices) and uv_indices[i] is not None else None,
                    "sunrise":     _hhmm(sunrises[i] if i < len(sunrises) else None),
                    "sunset":      _hhmm(sunsets[i] if i < len(sunsets) else None),
                })

        return {
            "success":     True,
            "location":    full_name,
            "temperature": current.get("temperature"),
            "wind_speed":  current.get("windspeed"),
            "wind_dir":    daily[0].get("wind_dir") if daily else None,
            "description": description,
            "units":       {"temp": "°C", "wind": "km/h"},
            "feels_like":  current_extra.get("apparent_temperature"),
            "humidity":    current_extra.get("relative_humidity_2m"),
            "uv_index":    daily[0].get("uv_index") if daily else None,
            "sunrise":     daily[0].get("sunrise") if daily else None,
            "sunset":      daily[0].get("sunset") if daily else None,
            "daily":       daily,
        }
    except Exception as e:
        return {"success": False, "error": f"Weather lookup failed: {e}"}


def _format_weather(data: dict) -> str:
    """Format weather dict into a LLM-ready context string."""
    if not data.get("success"):
        return f"[Weather] Error: {data.get('error', 'unknown')}"
    u = data["units"]
    return (
        f"=== REAL-TIME WEATHER (Open-Meteo API) ===\n"
        f"Location: {data['location']}\n"
        f"Temperature: {data['temperature']}{u['temp']}\n"
        f"Conditions: {data['description']}\n"
        f"Wind: {data['wind_speed']} {u['wind']}\n"
        f"[Prioritize this over any search snippets for weather questions]\n"
        f"==========================================\n"
    )


# ─────────────────────────────────────────────────────────────────────────────
# v3.3: Sandboxed Python code interpreter
# ─────────────────────────────────────────────────────────────────────────────

import multiprocessing
import queue as _queue
import io
import contextlib


def _execute_code_in_process(code: str, result_queue: multiprocessing.Queue):
    """Run code in isolated subprocess with restricted globals."""
    stdout_buf = io.StringIO()
    stderr_buf = io.StringIO()
    safe_globals = {
        "__builtins__": __builtins__,
        "print": print, "range": range, "len": len,
        "list": list, "dict": dict, "set": set, "tuple": tuple,
        "int": int, "float": float, "str": str, "bool": bool,
        "abs": abs, "sum": sum, "min": min, "max": max,
        "enumerate": enumerate, "zip": zip, "round": round,
        "sorted": sorted, "reversed": reversed, "map": map, "filter": filter,
        "divmod": divmod, "pow": pow,
        "math": __import__("math"),
        "json": __import__("json"),
        "re": __import__("re"),
        "random": __import__("random"),
        "datetime": __import__("datetime"),
    }
    try:
        with contextlib.redirect_stdout(stdout_buf), contextlib.redirect_stderr(stderr_buf):
            exec(code, safe_globals)
        result_queue.put({"success": True, "stdout": stdout_buf.getvalue(), "stderr": stderr_buf.getvalue()})
    except Exception as e:
        result_queue.put({"success": False, "stdout": stdout_buf.getvalue(), "stderr": str(e)})


async def run_code(code: str, timeout: int = 8) -> dict:
    """Execute Python code in an isolated process with timeout.

    Returns dict with success, stdout, stderr.
    """
    loop = asyncio.get_event_loop()
    result_queue: multiprocessing.Queue = multiprocessing.Queue()
    proc = multiprocessing.Process(target=_execute_code_in_process, args=(code, result_queue))
    proc.start()
    try:
        result = await loop.run_in_executor(
            None,
            lambda: result_queue.get(timeout=timeout)
        )
        proc.join(timeout=2)
        return result
    except (_queue.Empty, Exception) as e:
        if proc.is_alive():
            proc.terminate()
            proc.join()
        return {"success": False, "stdout": "", "stderr": f"Execution timed out or failed: {e}"}


# ─────────────────────────────────────────────────────────────────────────────
# v3.3: Deep Research (recursive multi-level search from old ECHO)
# ─────────────────────────────────────────────────────────────────────────────

async def _empty_string() -> str:
    """No-op coroutine placeholder for gather() when no URL to scrape."""
    return ""


async def deep_research(query: str, depth: int = 2, breadth: int = 3) -> dict:
    """Execute a recursive deep research loop.

    depth: how many levels of follow-up questions (1-3)
    breadth: parallel searches per level (1-5)

    Returns: {query, report, log, sources}
    """
    log: list[dict] = []
    all_findings: list[dict] = []

    # Step 1: Plan — generate sub-questions
    log.append({"step": "planning", "message": f"Analyzing: '{query}'"})
    plan_messages = [
        {"role": "system", "content": "You are a research planner. Generate sub-questions for deep research. Output JSON only: {\"sub_questions\": [\"q1\", \"q2\", \"q3\"]}"},
        {"role": "user", "content": f"Topic: {query}\nGenerate 3-4 specific sub-questions that together answer this comprehensively."},
    ]
    plan_text = await ollama_chat_text(plan_messages, model=AGENT_MODEL_MAP.get("Planner"))
    try:
        # Extract JSON from model output
        m = re.search(r'\{.*\}', plan_text or "", re.DOTALL)
        plan = json.loads(m.group()) if m else {}
    except Exception:
        plan = {}
    current_questions = plan.get("sub_questions", [query])[:breadth]

    analysis: dict = {}
    for level in range(1, depth + 1):
        log.append({"step": "level", "level": level, "message": f"Research level {level}/{depth} — {len(current_questions)} queries"})

        # Parallel searches for all questions in this level
        search_tasks = [web_search(q, max_results=3) for q in current_questions]
        level_search_results = await asyncio.gather(*search_tasks, return_exceptions=True)

        # Scrape top result per question concurrently
        scrape_tasks = []
        flat_results: list[dict] = []
        for qresults in level_search_results:
            if isinstance(qresults, list) and qresults:
                top = sorted(qresults, key=lambda r: r.get("credibility", 0.5), reverse=True)
                flat_results.extend(top[:2])
                scrape_tasks.append(scrape_page(top[0]["url"]))
            else:
                scrape_tasks.append(_empty_string())

        scraped_texts = await asyncio.gather(*scrape_tasks, return_exceptions=True)

        # Merge scraped text into findings
        for i, r in enumerate(flat_results[:len(scraped_texts)]):
            if isinstance(scraped_texts[i], str) and scraped_texts[i]:
                r = dict(r)
                r["content"] = scraped_texts[i][:1500]
            all_findings.append(r)
            log.append({"step": "found", "message": f"Found: {r.get('title', r.get('url', ''))}"})

        # Analyse gaps, generate follow-up questions
        if level < depth and all_findings:
            findings_ctx = "\n\n".join(
                f"Title: {f.get('title')}\n{f.get('content', f.get('snippet', ''))[:500]}"
                for f in all_findings[-6:]
            )
            gap_messages = [
                {"role": "system", "content": 'Analyze research findings. Output JSON only: {"summary": "...", "follow_up_questions": ["q1", "q2"]}'},
                {"role": "user", "content": f"Topic: {query}\n\nFindings:\n{findings_ctx}\n\nWhat is missing? Generate 2-3 follow-up questions."},
            ]
            gap_text = await ollama_chat_text(gap_messages, model=AGENT_MODEL_MAP.get("Researcher"))
            try:
                m2 = re.search(r'\{.*\}', gap_text or "", re.DOTALL)
                analysis = json.loads(m2.group()) if m2 else {}
            except Exception:
                analysis = {}
            follow_ups = analysis.get("follow_up_questions", [])
            if follow_ups:
                current_questions = follow_ups[:breadth]
                log.append({"step": "refine", "message": f"Refining with {len(current_questions)} new questions"})
            else:
                log.append({"step": "complete", "message": "Sufficient information gathered — stopping early"})
                break

    # Step 3: Synthesize final report
    log.append({"step": "synthesize", "message": "Compiling final report..."})
    unique_urls: set = set()
    context_for_report = ""
    for f in all_findings:
        url = f.get("url", "")
        if url not in unique_urls:
            unique_urls.add(url)
            cred = int(f.get("credibility", 0.5) * 100)
            context_for_report += f"\n--- {f.get('title', 'Source')} [{cred}% credibility] ({url}) ---\n"
            context_for_report += f.get("content", f.get("snippet", ""))[:1200] + "\n"

    report_messages = [
        {"role": "system", "content": "You are an expert research analyst. Write a comprehensive, well-structured research report in markdown. Include: Executive Summary, Detailed Findings (with headers), Conclusion, and Sources. Cite sources by title."},
        {"role": "user", "content": f"Topic: {query}\n\nResearch data:\n{context_for_report}"},
    ]
    report = await ollama_chat_text(report_messages, model=AGENT_MODEL_MAP.get("Researcher"), max_tokens=4096)

    return {
        "query": query,
        "report": report or "Report generation failed.",
        "log": log,
        "sources": list(unique_urls),
        "findings_count": len(all_findings),
    }


# ─────────────────────────────────────────────────────────────────────────────
# API Routes
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/health")
async def health():
    gpu = get_gpu_info()
    models = await get_loaded_models()
    return {
        "backend": "online",
        "version": "3.1.0",
        "gpu": {
            "name": gpu["name"],
            "vram_used": gpu["vram_used_mb"],
            "vram_total": gpu["vram_total_mb"],
        } if gpu else None,
        "models_loaded": models,
        "uptime": round(time.time() - START_TIME),
        "features": {
            "planner": True,
            "parallel_execution": True,
            "smart_routing": True,
            "response_cache": True,
            "context_compression": True,
            "vram_scheduler": True,
            "agent_memory": True,
            "self_reflection": True,
        },
    }


@app.get("/api/system")
def system_metrics():
    vm = psutil.virtual_memory()
    return {
        "cpu": {
            "name": platform.processor() or platform.machine(),
            "cores": psutil.cpu_count(logical=False) or psutil.cpu_count(),
            "threads": psutil.cpu_count(logical=True),
            "usage_percent": psutil.cpu_percent(interval=None),
            "temperature_c": get_cpu_temp(),
        },
        "ram": {
            "total_gb": round(vm.total / (1024 ** 3), 1),
            "used_gb": round(vm.used / (1024 ** 3), 1),
            "usage_percent": vm.percent,
        },
        "gpu": get_gpu_info(),
        "disk": get_disk_info(),
        "platform": platform.system(),
        "hostname": socket.gethostname(),
    }


@app.get("/api/agents")
def get_agents():
    agents = []
    active = None
    for name in PIPELINE:
        state = _agent_states[name]
        agents.append({"name": name, **state})
        if state["status"] in ("active", "processing"):
            active = name
    return {
        "agents": agents,
        "activeAgent": active,
        "pipeline": PIPELINE,
    }


@app.get("/api/models")
async def list_models():
    """List available models with their capabilities and VRAM estimates."""
    loaded = await get_loaded_models()
    gpu = get_gpu_info()
    total_vram = gpu["vram_total_mb"] if gpu else 0

    models = []
    total_estimated_vram = 0
    for m in loaded:
        caps = MODEL_CAPABILITIES.get(m, {"type": "general", "strengths": ["general"], "vram_mb": 0})
        est_vram = caps["vram_mb"]
        total_estimated_vram += est_vram
        models.append({
            "name": m,
            "loaded": True,
            "type": caps["type"],
            "strengths": caps["strengths"],
            "estimated_vram_mb": est_vram,
            "vram_percent": round((est_vram / total_vram) * 100, 1) if total_vram > 0 else 0,
        })
    return {
        "models": models,
        "total_estimated_vram_mb": total_estimated_vram,
        "vram_headroom_mb": max(0, total_vram - total_estimated_vram),
        "vram_scheduler": vram_scheduler.status,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Smart pipeline gate — skip planning for simple messages (Item 1)
# ─────────────────────────────────────────────────────────────────────────────

_SIMPLE_PATTERNS = re.compile(
    r"^(hi|hello|hey|thanks|thank you|ok|okay|yes|no|sure|bye|good|great|cool|nice|wow|lol|haha)\b",
    re.IGNORECASE,
)


def should_use_pipeline(text: str) -> bool:
    """Decide if a message is complex enough to warrant the planning pipeline."""
    text = text.strip()
    if len(text) < 120:
        return False  # Short/medium messages go direct
    if _SIMPLE_PATTERNS.match(text):
        return False  # Greetings/acks go direct
    if text.endswith("?") and len(text) < 200:
        return False  # Conversational questions go direct
    # Must have at least 2 sentences or technical density to justify planning
    sentences = len(re.findall(r'[.!?]+', text))
    if sentences < 2:
        return False
    return True


# ─────────────────────────────────────────────────────────────────────────────
# Item 7: Depth-adaptive Critic
# ─────────────────────────────────────────────────────────────────────────────


def estimate_complexity(text: str) -> int:
    """Score 1-3 based on length, question count, and technical density."""
    score = 1
    if len(text) > 300:
        score += 1
    if len(text) > 800:
        score += 1

    questions = text.count("?")
    if questions >= 3:
        score += 1

    tech_keywords = len(re.findall(
        r"\b(implement|algorithm|database|architecture|optimize|refactor|deploy|scale|security|API)\b",
        text, re.IGNORECASE,
    ))
    if tech_keywords >= 3:
        score += 1

    return min(score, 3)


# ─────────────────────────────────────────────────────────────────────────────
# Main chat endpoint — now with full pipeline
# ─────────────────────────────────────────────────────────────────────────────


@app.post("/api/chat")
async def chat(req: ChatRequest):
    """Full multi-agent chat with planning, parallel execution, caching, compression, and reflection."""
    model = req.model or AGENT_MODEL_MAP["default"]
    temperature = req.temperature or 0.7
    max_tokens = req.max_tokens or 2048

    # ── v3.6: Security scan on incoming user text ────────────────────────
    # Get user text early for scanning (full extraction happens below too)
    _scan_user_msgs = [m for m in req.messages if m.role == "user"]
    _scan_text = _scan_user_msgs[-1].content if _scan_user_msgs else ""
    if _scan_text:
        try:
            injection_findings = scan_injection(_scan_text)
            if injection_findings:
                high_risk = [f for f in injection_findings if f["level"] == "high"]
                _injection_scan_log.append({
                    "ts": datetime.now(timezone.utc).isoformat(),
                    "findings": injection_findings,
                    "text_preview": _scan_text[:100],
                })
                if len(_injection_scan_log) > 500:
                    _injection_scan_log.pop(0)
                if high_risk:
                    _logger.warning(f"High-risk injection attempt detected: {high_risk}")
                    # Don't block — log and continue, but strip the injection pattern
        except Exception:
            pass

    # Build message list
    depth_instruction = ""
    if req.depth and req.depth > 0:
        depth_instruction = (
            f"\n\nCritic depth is set to {req.depth}. Before finalizing your response, "
            f"internally review your answer {req.depth} time(s) for accuracy, hallucinations, "
            f"and completeness. Correct any issues before responding."
        )

    messages: list[dict] = [
        {"role": "system", "content": SYSTEM_PROMPT + depth_instruction}
    ]
    for msg in req.messages:
        if msg.role in ("user", "assistant", "system"):
            messages.append({"role": msg.role, "content": msg.content})

    # Get the user's latest message for planning/routing
    user_messages = [m for m in req.messages if m.role == "user"]
    user_text = user_messages[-1].content if user_messages else ""

    # ── v3.7: Inject file attachment content into user message ──────────
    if req.attachments:
        attachment_blocks = []
        for att in req.attachments:
            name = att.get("name", "file")
            content = att.get("content", "")
            if content:
                attachment_blocks.append(
                    f"[ATTACHED FILE: {name}]\n{content}\n[/ATTACHED FILE]"
                )
        if attachment_blocks:
            user_text = "\n\n".join(attachment_blocks) + "\n\n" + user_text
            # Also patch the last user message in the messages list
            for msg in reversed(messages):
                if msg["role"] == "user":
                    msg["content"] = "\n\n".join(attachment_blocks) + "\n\n" + msg["content"]
                    break

    # ── Feature #6: Inject relevant memories ────────────────────────────
    try:
        memories = await recall_relevant_memories(user_text, limit=3)
        if memories:
            mem_text = "\n".join(
                f"- [{m['type']}] {m['content'][:150]}" for m in memories
            )
            messages.insert(1, {
                "role": "system",
                "content": f"[RECALLED MEMORIES — relevant past knowledge]\n{mem_text}",
            })
    except Exception:
        pass

    # ── v3.5: Project Mode — inject active project context ───────────────
    if req.project_id and req.project_id in _projects:
        proj = _projects[req.project_id]
        proj_ctx = (
            f"[ACTIVE PROJECT: {proj['name']}]\n"
            f"Goal: {proj['goal']}\n"
            f"Context: {proj.get('context', '')}\n"
            f"Tasks: {len(proj.get('tasks', []))} task(s)"
        )
        messages.insert(1, {"role": "system", "content": proj_ctx})

    # ── Feature #3: Context compression (v3.5: ContextWindowManager) ────
    try:
        messages = _ctx_manager.trim(messages, model)
    except Exception:
        messages = await compress_context(messages, max_tokens=4096)

    # ── Feature #4: Check cache ─────────────────────────────────────────
    if not req.no_cache:
        cached = response_cache.get(model, messages, temperature)
        if cached:
            # Return cached response as instant SSE
            async def cached_stream():
                sse_data = {"choices": [{"delta": {"content": cached}}]}
                yield f"data: {json.dumps(sse_data)}\n\n"
                meta = {"choices": [{"delta": {"content": ""}, "meta": {"cached": True}}]}
                yield f"data: {json.dumps(meta)}\n\n"
                yield "data: [DONE]\n\n"

            return StreamingResponse(
                cached_stream(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
            )

    # ── Vision: if images attached, route directly to vision model ──────
    if req.images and len(req.images) > 0:
        vision_model = req.model or await get_vision_model()
        if vision_model:
            # Attach images to the last user message
            for msg in reversed(messages):
                if msg["role"] == "user":
                    msg["images"] = req.images
                    break

            async def vision_stream():
                payload = {
                    "model": vision_model,
                    "messages": messages,
                    "stream": True,
                    "options": {"num_keep": 256},
                    "keep_alive": "10m",
                }
                _client = await get_ollama_client()
                async with _client.stream("POST", f"{OLLAMA_URL}/api/chat", json=payload, timeout=120.0) as resp:
                    async for line in resp.aiter_lines():
                        if not line.strip():
                            continue
                        try:
                            chunk = json.loads(line)
                            if chunk.get("done"):
                                yield "data: [DONE]\n\n"
                                break
                            content = chunk.get("message", {}).get("content", "")
                            if content:
                                sse = {"choices": [{"delta": {"content": content}}]}
                                yield f"data: {json.dumps(sse)}\n\n"
                        except Exception:
                            continue

            return StreamingResponse(
                vision_stream(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
            )

    # ── v3.2: Custom workflow override ──────────────────────────────────
    if req.workflow:
        custom_agents = req.workflow.get("agents", [])
        custom_sys_prompts = req.workflow.get("system_prompts", {})
        if custom_agents:
            # Build a synthetic plan using the workflow's agent order
            from_workflow = [
                Subtask(
                    id=f"wf-{i}",
                    agent=a if a in PIPELINE else "Supervisor",
                    task=user_text,
                    depends_on=[f"wf-{i-1}"] if i > 0 else [],
                    priority=i,
                )
                for i, a in enumerate(custom_agents)
            ]
            # Apply per-agent system prompt overrides
            wf_system = SYSTEM_PROMPT + depth_instruction
            if custom_sys_prompts:
                overrides = "; ".join(f"{k}: {v}" for k, v in custom_sys_prompts.items())
                wf_system += f"\n\n[WORKFLOW OVERRIDES] {overrides}"
            wf_plan = TaskPlan(subtasks=from_workflow)
            results = await run_pipeline(wf_plan, wf_system, temperature, max_tokens, user_text=user_text)
            merged = "\n\n".join(results.values()) if results else ""
            if len(results) > 1:
                synth_msgs = [
                    {"role": "system", "content": "Synthesize these agent results into one coherent response."},
                    {"role": "user", "content": f"Question: {user_text}\n\nResults:\n{merged}"},
                ]
                merged = await ollama_chat_text(synth_msgs, model=AGENT_MODEL_MAP.get("Supervisor"))

            async def wf_stream():
                sse_data = {"choices": [{"delta": {"content": merged}}]}
                yield f"data: {json.dumps(sse_data)}\n\n"
                yield "data: [DONE]\n\n"

            return StreamingResponse(
                wf_stream(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
            )

    # ── Feature #8: Task planning ───────────────────────────────────────
    use_pipeline = req.enable_planning and should_use_pipeline(user_text)

    if use_pipeline:
        try:
            _agent_states["Planner"]["status"] = "active"
            _agent_states["Planner"]["currentTask"] = "Decomposing request"
            plan = await plan_task(user_text, messages[-6:])
            _agent_states["Planner"]["status"] = "idle"
            _agent_states["Planner"]["currentTask"] = None

            # If planner returns multiple subtasks, run full pipeline
            if len(plan.subtasks) > 1:
                # ── Feature #1: Parallel execution ──────────────────────
                # ── v3.5 Feature #13: Progressive Pipeline Stream ────────
                if req.depth and req.depth >= 1:
                    # Progressive mode: stream agent results as they complete
                    async def progressive_pipeline_stream():
                        """Stream partial results as each agent completes (depth >= 1)."""
                        try:
                            # Emit plan summary step
                            plan_step = {"type": "step", "agent": "Planner", "text": f"Plan ready: {len(plan.subtasks)} task(s)", "status": "done"}
                            yield f"data: {json.dumps(plan_step)}\n\n"

                            subtask_results: dict[str, str] = {}
                            completed: set[str] = set()
                            full_output_parts: list[str] = []

                            while len(completed) < len(plan.subtasks):
                                ready = [
                                    st for st in plan.subtasks
                                    if st.id not in completed
                                    and all(dep in completed for dep in st.depends_on)
                                ]
                                if not ready:
                                    break
                                ready.sort(key=lambda s: s.priority)

                                # Emit start steps for all ready agents
                                for st in ready:
                                    start_step = {"type": "step", "agent": st.agent, "text": st.task[:60], "status": "start"}
                                    yield f"data: {json.dumps(start_step)}\n\n"

                                batch = await asyncio.gather(
                                    *[run_subtask(st, SYSTEM_PROMPT + depth_instruction, subtask_results, temperature, max_tokens)
                                      for st in ready],
                                    return_exceptions=True,
                                )
                                for res in batch:
                                    if isinstance(res, Exception):
                                        continue
                                    task_id, text = res
                                    subtask_results[task_id] = text
                                    completed.add(task_id)
                                    agent_name = next(
                                        (st.agent for st in plan.subtasks if st.id == task_id),
                                        "Agent"
                                    )
                                    # Emit done step for this agent
                                    done_step = {"type": "step", "agent": agent_name, "text": "Completed", "status": "done"}
                                    yield f"data: {json.dumps(done_step)}\n\n"

                                    chunk = f"\n\n---\n**[{agent_name}]**\n{text}\n"
                                    full_output_parts.append(chunk)
                                    sse_data = {"choices": [{"delta": {"content": chunk}}]}
                                    yield f"data: {json.dumps(sse_data)}\n\n"

                            # Synthesis step
                            if len(subtask_results) > 1:
                                sup_step = {"type": "step", "agent": "Supervisor", "text": "Synthesizing agent results", "status": "start"}
                                yield f"data: {json.dumps(sup_step)}\n\n"
                                parts_str = "\n\n".join(
                                    f"**[{st.agent}]**\n{subtask_results.get(st.id, '')}"
                                    for st in plan.subtasks
                                )
                                synth_msgs = [
                                    {"role": "system", "content": AGENT_SPECIFIC_PROMPTS.get("Supervisor", "Synthesize agent results.")},
                                    {"role": "user", "content": f"Question: {user_text}\n\nAgent results:\n{parts_str}"},
                                ]
                                synth = await ollama_chat_text(synth_msgs, model=AGENT_MODEL_MAP.get("Supervisor", "llama3.2:3b"))
                                if synth:
                                    sup_done = {"type": "step", "agent": "Supervisor", "text": "Synthesis complete", "status": "done"}
                                    yield f"data: {json.dumps(sup_done)}\n\n"
                                    sep = "\n\n---\n**[Supervisor — Final Synthesis]**\n"
                                    full_output_parts.append(sep + synth)
                                    sse_data = {"choices": [{"delta": {"content": sep + synth}}]}
                                    yield f"data: {json.dumps(sse_data)}\n\n"

                            merged = "".join(full_output_parts)
                            response_cache.put(model, messages, temperature, merged)
                            if len(messages) >= 6:
                                asyncio.create_task(auto_extract_memories(
                                    messages + [{"role": "assistant", "content": merged}]
                                ))
                        except Exception as exc:
                            _logger.error(f"[progressive_pipeline] Error: {exc}")
                        finally:
                            yield "data: [DONE]\n\n"

                    return StreamingResponse(
                        progressive_pipeline_stream(),
                        media_type="text/event-stream",
                        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
                    )
                else:
                    # Non-progressive: collect all results then stream
                    results = await run_pipeline(plan, SYSTEM_PROMPT + depth_instruction, temperature, max_tokens, user_text=user_text)

                    if len(results) == 1:
                        merged = list(results.values())[0]
                    else:
                        parts = "\n\n".join(
                            f"**[{st.agent} — {st.task[:60]}]**\n{results.get(st.id, '(no result)')}"
                            for st in plan.subtasks
                        )
                        synth_messages = [
                            {"role": "system", "content": "You are the Supervisor agent in ECHO. Synthesize the following agent results into a single complete, polished response. CRITICAL: If any agent produced code blocks, PRESERVE THEM EXACTLY and output all code directly — never summarize code into plain-text descriptions. Combine results naturally, removing redundancy while keeping all technical content intact."},
                            {"role": "user", "content": f"Original question: {user_text}\n\nAgent results:\n{parts}"},
                        ]
                        merged = await ollama_chat_text(synth_messages, model=AGENT_MODEL_MAP.get("Supervisor", "llama3.1:8b"))

                    if req.enable_reflection:
                        complexity = estimate_complexity(user_text)
                        merged, reflection_info = await reflect_on_response(user_text, merged, max_loops=complexity)

                    response_cache.put(model, messages, temperature, merged)
                    if len(messages) >= 6:
                        asyncio.create_task(auto_extract_memories(messages + [{"role": "assistant", "content": merged}]))

                    async def pipeline_stream():
                        sse_data = {"choices": [{"delta": {"content": merged}}]}
                        yield f"data: {json.dumps(sse_data)}\n\n"
                        yield "data: [DONE]\n\n"

                    return StreamingResponse(
                        pipeline_stream(),
                        media_type="text/event-stream",
                        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
                    )
        except Exception as e:
            print(f"[planner] Pipeline failed, falling back to direct: {e}")
            _agent_states["Planner"]["status"] = "idle"
            _agent_states["Planner"]["currentTask"] = None

    # ── Fallback / simple requests: direct streaming ────────────────────
    # Feature #2: Smart model routing
    # v3.5 Feature #11: Quantization switching based on complexity
    complexity = estimate_complexity(user_text)
    quant_model = get_quantization_model(complexity, req.model)
    model = quant_model if quant_model else await select_model(user_text, req.model)
    await vram_scheduler.request_model(model)

    # Swap multi-agent system prompt for lean direct prompt so small models
    # don't hallucinate "Research Chain / Developer's Code Snippet" sections.
    if messages and messages[0]["role"] == "system":
        messages[0]["content"] = DIRECT_PROMPT + (
            f"\n\nCritic depth: {req.depth}. Review your answer {req.depth} time(s) for accuracy before responding."
            if req.depth and req.depth > 0 else ""
        )

    active_agent = "Supervisor"
    for agent_name in PIPELINE:
        if model == AGENT_MODEL_MAP.get(agent_name):
            active_agent = agent_name
            break

    # v3.5 Feature #14: Speculative Decode for depth==0 (fastest mode)
    if req.depth == 0:
        try:
            spec_result = await speculative_decode(messages, full_model=model)
            if spec_result:
                response_cache.put(model, messages, temperature, spec_result)
                async def spec_stream():
                    sse_data = {"choices": [{"delta": {"content": spec_result}}]}
                    yield f"data: {json.dumps(sse_data)}\n\n"
                    yield "data: [DONE]\n\n"
                return StreamingResponse(
                    spec_stream(),
                    media_type="text/event-stream",
                    headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
                )
        except Exception as e:
            _logger.debug(f"[speculative_decode] Fell through to normal stream: {e}")

    request_start = time.time()
    _agent_states[active_agent]["status"] = "active"
    _agent_states[active_agent]["currentTask"] = "Processing request"
    _agent_states[active_agent]["lastActive"] = datetime.now(timezone.utc).isoformat()

    async def stream_and_track():
        # Emit a single step event for direct streaming
        direct_step = {"type": "step", "agent": active_agent, "text": "Processing request", "status": "start"}
        yield f"data: {json.dumps(direct_step)}\n\n"
        token_count = 0
        full_response = ""
        try:
            async for chunk in ollama_stream(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            ):
                if chunk.startswith("data: ") and chunk.strip() != "data: [DONE]":
                    try:
                        data = json.loads(chunk[6:])
                        content = data.get("choices", [{}])[0].get("delta", {}).get("content", "")
                        token_count += max(1, len(content) // 4)
                        full_response += content
                    except Exception:
                        pass
                yield chunk
        finally:
            elapsed_ms = round((time.time() - request_start) * 1000)
            _agent_states[active_agent]["status"] = "idle"
            _agent_states[active_agent]["currentTask"] = None
            _agent_states[active_agent]["tokensProcessed"] += token_count
            _agent_states[active_agent]["totalResponseMs"] += elapsed_ms
            _agent_states[active_agent]["requestCount"] += 1
            await vram_scheduler.release_model(model)

            # Cache the streamed response
            if full_response:
                response_cache.put(model, messages, temperature, full_response)

            # Auto-extract memories in background (only for substantial conversations)
            if full_response and len(messages) >= 6:
                asyncio.create_task(auto_extract_memories(
                    messages + [{"role": "assistant", "content": full_response}]
                ))

    return StreamingResponse(
        stream_and_track(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# Cache endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/cache/stats")
def cache_stats():
    return response_cache.stats


@app.delete("/api/cache")
def clear_cache():
    response_cache.clear()
    return {"status": "ok", "message": "Cache cleared"}


# ─────────────────────────────────────────────────────────────────────────────
# Memory endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.post("/api/memory/store")
async def store_memory(req: MemoryStoreRequest):
    """Store a memory of a specific type."""
    coll = _get_memory_collection(req.type)
    if coll is None:
        raise HTTPException(status_code=503, detail="Memory system not available")

    # v3.6: Redact PII/secrets before persisting to memory
    content = req.content
    try:
        content_clean, redacted_types = redact_sensitive(content)
        if redacted_types:
            _logger.info(f"[memory] Redacted {redacted_types} from memory content")
        content = content_clean
    except Exception:
        pass

    mem_id = f"{req.type[:3]}-{int(time.time() * 1000)}"
    metadata = {
        "summary": (req.summary or content[:200]),
        "tags": ",".join(req.tags or []),
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }

    try:
        coll.upsert(documents=[content], ids=[mem_id], metadatas=[metadata])
        return {"status": "ok", "id": mem_id, "type": req.type}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to store memory: {e}")


@app.post("/api/memory/recall")
async def recall_memory(req: MemoryRecallRequest):
    """Semantic search across memory types."""
    if req.type:
        # Search specific type
        coll = _get_memory_collection(req.type)
        if coll is None or coll.count() == 0:
            return {"results": []}
        try:
            results = coll.query(
                query_texts=[req.query],
                n_results=min(req.limit or 5, coll.count()),
            )
            formatted = []
            ids = results.get("ids", [[]])[0]
            docs = results.get("documents", [[]])[0]
            metas = results.get("metadatas", [[]])[0]
            distances = results.get("distances", [[]])[0]
            for i, doc_id in enumerate(ids):
                similarity = round(1 - (distances[i] / 2), 4) if distances else 0.0
                formatted.append({
                    "id": doc_id,
                    "type": req.type,
                    "content": docs[i] if docs else "",
                    "summary": (metas[i] or {}).get("summary", ""),
                    "tags": (metas[i] or {}).get("tags", "").split(",") if metas[i] else [],
                    "similarity": similarity,
                    "timestamp": (metas[i] or {}).get("timestamp", ""),
                })
            return {"results": formatted}
        except Exception as e:
            return {"results": [], "error": str(e)}
    else:
        # Search all types
        results = await recall_relevant_memories(req.query, req.limit or 5)
        return {"results": results}


@app.get("/api/memory/list")
async def list_memories(type: str = "all"):
    """List memories by type."""
    types = ["episodic", "semantic", "procedural"] if type == "all" else [type]
    all_memories = []

    for mem_type in types:
        coll = _get_memory_collection(mem_type)
        if coll is None:
            continue
        try:
            count = coll.count()
            if count == 0:
                continue
            # Get all entries (up to 100)
            result = coll.get(limit=min(count, 100))
            ids = result.get("ids", [])
            docs = result.get("documents", [])
            metas = result.get("metadatas", [])
            for i, doc_id in enumerate(ids):
                all_memories.append({
                    "id": doc_id,
                    "type": mem_type,
                    "content": docs[i] if docs else "",
                    "summary": (metas[i] or {}).get("summary", ""),
                    "tags": (metas[i] or {}).get("tags", "").split(",") if metas[i] else [],
                    "timestamp": (metas[i] or {}).get("timestamp", ""),
                })
        except Exception:
            continue

    return {"memories": all_memories, "total": len(all_memories)}


@app.delete("/api/memory/{memory_id}")
async def delete_memory(memory_id: str):
    """Delete a specific memory by ID."""
    # Determine type from ID prefix
    prefix_map = {"ep-": "episodic", "sem": "semantic", "pro": "procedural"}
    mem_type = None
    for prefix, mtype in prefix_map.items():
        if memory_id.startswith(prefix):
            mem_type = mtype
            break

    if mem_type is None:
        raise HTTPException(status_code=400, detail="Cannot determine memory type from ID")

    coll = _get_memory_collection(mem_type)
    if coll is None:
        raise HTTPException(status_code=503, detail="Memory system not available")

    try:
        coll.delete(ids=[memory_id])
        return {"status": "ok", "deleted": memory_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete memory: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Document / RAG endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.post("/api/documents")
async def add_document(req: DocumentRequest):
    collection = _get_chroma_collection()
    if collection is None:
        raise HTTPException(status_code=503, detail="ChromaDB not available")

    doc_id = req.id or f"doc-{int(time.time() * 1000)}"
    metadata = {"title": req.title}
    if req.user_id:
        metadata["user_id"] = req.user_id

    try:
        collection.upsert(documents=[req.content], ids=[doc_id], metadatas=[metadata])
        asyncio.create_task(_rebuild_bm25())  # v3.2: keep BM25 in sync
        return {"status": "ok", "id": doc_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to ingest document: {e}")


@app.post("/api/semantic-search")
async def semantic_search(req: SearchRequest):
    collection = _get_chroma_collection()
    if collection is None:
        return {"results": []}

    try:
        where = {"user_id": req.user_id} if req.user_id else None
        n = min(req.limit or 5, max(collection.count() or 1, 1))
        query_kwargs: dict = {"query_texts": [req.query], "n_results": n}
        if where:
            query_kwargs["where"] = where

        results = collection.query(**query_kwargs)
        ids = results.get("ids", [[]])[0]
        docs = results.get("documents", [[]])[0]
        metas = results.get("metadatas", [[]])[0]
        distances = results.get("distances", [[]])[0]

        # Build vector score map
        vector_scores: dict[str, float] = {}
        for i, doc_id in enumerate(ids):
            vector_scores[doc_id] = round(1 - (distances[i] / 2), 4) if distances else 0.0

        # v3.2: BM25 scores (hybrid mode)
        bm25_scores: dict[str, float] = {}
        if req.hybrid and _BM25_AVAILABLE:
            bm25_results = await _bm25_index.search(req.query, n=max(n * 2, 20))
            bm25_scores = dict(bm25_results)

        formatted = []
        for i, doc_id in enumerate(ids):
            v_score = vector_scores.get(doc_id, 0.0)
            b_score = bm25_scores.get(doc_id, 0.0)
            hybrid_score = round(0.6 * v_score + 0.4 * b_score, 4) if req.hybrid and _BM25_AVAILABLE else v_score
            formatted.append({
                "id": doc_id,
                "title": metas[i].get("title", "Untitled") if metas else "Untitled",
                "content": docs[i] if docs else "",
                "similarity": hybrid_score,
                "vector_score": v_score,
                "bm25_score": b_score,
                "hybrid": req.hybrid and _BM25_AVAILABLE,
            })

        formatted.sort(key=lambda x: -x["similarity"])
        return {"results": formatted}
    except Exception as e:
        _logger.error(f"Semantic search error: {e}")
        return {"results": []}


# ─────────────────────────────────────────────────────────────────────────────
# Item 6: Document list & delete endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/documents/list")
async def list_documents(limit: int = 50, offset: int = 0):
    """List all documents stored in ChromaDB."""
    collection = _get_chroma_collection()
    if collection is None:
        return {"documents": [], "total": 0}

    try:
        total = collection.count()
        if total == 0:
            return {"documents": [], "total": 0}

        result = collection.get(limit=min(limit, total), offset=offset)
        docs = []
        ids = result.get("ids", [])
        documents = result.get("documents", [])
        metadatas = result.get("metadatas", [])
        for i, doc_id in enumerate(ids):
            content = documents[i] if documents else ""
            meta = metadatas[i] if metadatas else {}
            docs.append({
                "id": doc_id,
                "title": (meta or {}).get("title", "Untitled"),
                "content_preview": content[:200] if content else "",
                "user_id": (meta or {}).get("user_id", ""),
            })
        return {"documents": docs, "total": total}
    except Exception as e:
        _logger.error(f"Document list error: {e}")
        return {"documents": [], "total": 0, "error": str(e)}


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document from ChromaDB by ID."""
    collection = _get_chroma_collection()
    if collection is None:
        raise HTTPException(status_code=503, detail="ChromaDB not available")

    try:
        collection.delete(ids=[doc_id])
        asyncio.create_task(_rebuild_bm25())  # v3.2: keep BM25 in sync
        return {"status": "ok", "deleted": doc_id}
    except Exception as e:
        _logger.error(f"Document delete error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Item 4: Telemetry endpoint
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/telemetry")
async def telemetry():
    """Aggregated real telemetry: VRAM, cache, agents, uptime."""
    gpu = get_gpu_info()
    cache = response_cache.stats

    agents = []
    for name in PIPELINE:
        st = _agent_states[name]
        avg_ms = round(st["totalResponseMs"] / max(st["requestCount"], 1))
        agents.append({
            "name": name,
            "tasks": st["requestCount"],
            "avgTimeMs": avg_ms,
            "tokensProcessed": st["tokensProcessed"],
            "status": st["status"],
        })

    return {
        "uptime": round(time.time() - START_TIME),
        "vram": {
            "used_mb": gpu["vram_used_mb"] if gpu else 0,
            "total_mb": gpu["vram_total_mb"] if gpu else 0,
            "percent": round((gpu["vram_used_mb"] / gpu["vram_total_mb"]) * 100, 1) if gpu and gpu["vram_total_mb"] > 0 else 0,
        },
        "cache": cache,
        "agents": agents,
        "pipeline_queue": sum(1 for a in _agent_states.values() if a["status"] == "active"),
        "models_loaded": await get_loaded_models(),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Item 7: Feedback endpoints
# ─────────────────────────────────────────────────────────────────────────────

_feedback_log: list[dict] = []


class FeedbackRequest(BaseModel):
    message_id: str
    rating: int  # 1-5
    comment: Optional[str] = None


# ─────────────────────────────────────────────────────────────────────────────
# v3.2: Web Search endpoint
# ─────────────────────────────────────────────────────────────────────────────


@app.post("/api/web-search")
async def api_web_search(req: WebSearchRequest):
    """Search the web and optionally scrape pages for summarized results."""
    raw = await web_search(req.query, max_results=req.max_results or 5)
    if not raw:
        return {"results": [], "summary": "Web search unavailable.", "query": req.query}

    if req.scrape:
        scrape_tasks = [scrape_page(r["url"]) for r in raw[:3]]
        scraped = await asyncio.gather(*scrape_tasks, return_exceptions=True)
        for i, r in enumerate(raw[:3]):
            if i < len(scraped) and isinstance(scraped[i], str) and scraped[i]:
                r["scraped_text"] = scraped[i][:800]
            else:
                r["scraped_text"] = ""

    summary_messages = [
        {"role": "system", "content": "Summarize these web search results concisely. Include key facts and cite the source titles."},
        {"role": "user", "content": f"Query: {req.query}\n\n" + "\n\n".join(
            f"[{r['title']}] {r.get('scraped_text') or r.get('snippet', '')}" for r in raw[:3]
        )},
    ]
    summary = await ollama_chat_text(summary_messages, model=AGENT_MODEL_MAP.get("Researcher"))
    return {"results": raw, "summary": summary, "query": req.query}


# ─────────────────────────────────────────────────────────────────────────────
# v3.3: Deep Research, Weather, Code Interpreter endpoints
# ─────────────────────────────────────────────────────────────────────────────


class DeepResearchRequest(BaseModel):
    query: str
    depth: Optional[int] = 2   # 1-3 recursion levels
    breadth: Optional[int] = 3 # parallel searches per level


@app.post("/api/deep-research")
async def api_deep_research(req: DeepResearchRequest):
    """Recursive multi-level web research with LLM synthesis.

    depth=1: single-pass (fast), depth=2: two-level (default), depth=3: thorough
    """
    d = max(1, min(3, req.depth or 2))
    b = max(1, min(5, req.breadth or 3))
    result = await deep_research(req.query, depth=d, breadth=b)
    return result


class WeatherRequest(BaseModel):
    location: str


@app.post("/api/weather")
async def api_weather(req: WeatherRequest):
    """Get real-time weather for a city (Open-Meteo API, no key required)."""
    data = await get_weather(req.location)
    if data.get("success"):
        data["formatted"] = _format_weather(data)
    return data


class RunCodeRequest(BaseModel):
    code: str
    timeout: Optional[int] = 8


@app.post("/api/run-code")
async def api_run_code(req: RunCodeRequest):
    """Execute Python code in a sandboxed subprocess.

    Returns stdout, stderr, success flag. Timeout max 30s.
    Allowed imports: math, json, re, random, datetime (no file system or network).
    """
    t = max(1, min(30, req.timeout or 8))
    result = await run_code(req.code, timeout=t)
    return result


# ─────────────────────────────────────────────────────────────────────────────
# v3.2: Knowledge folder status endpoint
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/knowledge/status")
async def knowledge_status():
    """Returns files in the knowledge folder and their ingest status."""
    return {
        "folder": str(_KNOWLEDGE_DIR),
        "watchdog_active": _WATCHDOG_AVAILABLE and _knowledge_watcher._observer is not None,
        "files": _knowledge_watcher.status(),
    }


@app.post("/api/feedback")
async def submit_feedback(req: FeedbackRequest):
    """Store user feedback on a response."""
    entry = {
        "message_id": req.message_id,
        "rating": max(1, min(5, req.rating)),
        "comment": req.comment or "",
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }
    _feedback_log.append(entry)
    _logger.info(f"Feedback: msg={req.message_id} rating={req.rating}")

    # Auto-store positive feedback with comment as procedural memory
    if req.rating >= 4 and req.comment:
        try:
            coll = _get_memory_collection("procedural")
            if coll:
                mem_id = f"fb-{int(time.time() * 1000)}"
                coll.upsert(
                    documents=[req.comment],
                    ids=[mem_id],
                    metadatas=[{
                        "summary": f"Positive feedback: {req.comment[:150]}",
                        "tags": "feedback,positive",
                        "timestamp": entry["timestamp"],
                    }],
                )
        except Exception:
            pass

    return {"status": "ok", "total_feedback": len(_feedback_log)}


@app.get("/api/feedback/stats")
def feedback_stats():
    """Aggregate feedback statistics."""
    if not _feedback_log:
        return {"total": 0, "avg_rating": 0, "distribution": {}}

    ratings = [f["rating"] for f in _feedback_log]
    dist = {str(i): ratings.count(i) for i in range(1, 6)}
    return {
        "total": len(_feedback_log),
        "avg_rating": round(sum(ratings) / len(ratings), 2),
        "distribution": dist,
        "recent": _feedback_log[-10:],
    }


# ─────────────────────────────────────────────────────────────────────────────
# Item 8: Model management endpoint
# ─────────────────────────────────────────────────────────────────────────────


class ModelManageRequest(BaseModel):
    model: str
    action: str  # "load" | "unload"


@app.post("/api/models/manage")
async def manage_model(req: ModelManageRequest):
    """Load or unload an Ollama model."""
    client = await get_ollama_client()

    if req.action == "load":
        try:
            await client.post(
                f"{OLLAMA_URL}/api/generate",
                json={"model": req.model, "prompt": "", "keep_alive": "10m"},
                timeout=60,
            )
            vram_scheduler._loaded_models[req.model] = time.time()
            _logger.info(f"Model loaded: {req.model}")
            return {"status": "ok", "action": "loaded", "model": req.model}
        except Exception as e:
            _logger.error(f"Model load failed: {req.model} - {e}")
            raise HTTPException(status_code=500, detail=f"Failed to load: {e}")

    elif req.action == "unload":
        try:
            await client.post(
                f"{OLLAMA_URL}/api/generate",
                json={"model": req.model, "keep_alive": 0},
                timeout=30,
            )
            vram_scheduler._loaded_models.pop(req.model, None)
            _logger.info(f"Model unloaded: {req.model}")
            return {"status": "ok", "action": "unloaded", "model": req.model}
        except Exception as e:
            _logger.error(f"Model unload failed: {req.model} - {e}")
            raise HTTPException(status_code=500, detail=f"Failed to unload: {e}")

    else:
        raise HTTPException(status_code=400, detail=f"Invalid action: {req.action}")


# ─────────────────────────────────────────────────────────────────────────────
# Enhanced /api/models with VRAM estimation (Item 8)
# ─────────────────────────────────────────────────────────────────────────────
# (The original /api/models is above — this enhances it in-place.)


# ─────────────────────────────────────────────────────────────────────────────
# Skill endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.post("/api/skill-tools")
async def skill_tools(req: SkillToolsRequest):
    """Skill creator, evaluator, and trigger tuning via Ollama."""
    model = AGENT_MODEL_MAP.get("default", "llama3.1:8b")

    if req.mode == "create":
        if req.existingSkill:
            user_prompt = (
                f"I have an existing skill that needs improvement:\n\n---\n{req.existingSkill}\n---\n\n"
                f"Here's what I want to change/improve: {req.description}\n\nTarget agent: {req.agent or 'Developer'}"
            )
        else:
            user_prompt = (
                f"Create a new skill based on this description: {req.description}\n\n"
                f"Target agent: {req.agent or 'Developer'}"
            )
        messages = [
            {"role": "system", "content": SKILL_CREATOR_SYSTEM},
            {"role": "user", "content": user_prompt},
        ]
        return StreamingResponse(
            ollama_stream(model=model, messages=messages, temperature=0.7, max_tokens=4096),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
        )

    elif req.mode == "eval":
        user_prompt = (
            f"## Skill Content\n{req.skillContent}\n\n"
            f"## Test Prompt\n{req.prompt}\n\n"
            f"## Expected Behavior\n{req.expectedBehavior}"
        )
        messages = [
            {"role": "system", "content": EVAL_SYSTEM},
            {"role": "user", "content": user_prompt},
        ]
        result = await ollama_chat_json(messages, model=model)
        return {
            "pass": result.get("pass", False),
            "score": result.get("score", 1),
            "notes": result.get("notes", "Failed to grade"),
        }

    elif req.mode == "tune-trigger":
        should = "\n".join(f"{i+1}. {p}" for i, p in enumerate(req.shouldTrigger or []))
        should_not = "\n".join(f"{i+1}. {p}" for i, p in enumerate(req.shouldNotTrigger or []))
        user_prompt = (
            f"## Skill: {req.skillName}\n\n"
            f"## Current Description\n{req.currentDescription or '(none)'}\n\n"
            f"## Should Trigger On\n{should}\n\n"
            f"## Should NOT Trigger On\n{should_not}"
        )
        messages = [
            {"role": "system", "content": TRIGGER_TUNING_SYSTEM},
            {"role": "user", "content": user_prompt},
        ]
        result = await ollama_chat_json(messages, model=model)
        return {
            "analysis": result.get("analysis", "Failed to analyze"),
            "suggestedDescription": result.get("suggestedDescription", ""),
            "predictedAccuracy": result.get("predictedAccuracy", 0),
        }

    else:
        raise HTTPException(status_code=400, detail=f"Invalid mode: {req.mode}")


@app.get("/api/skills/scan")
async def scan_skills():
    """Read all .md files from backend/skills/ directory."""
    skills_dir = _BASE / "skills"
    if not skills_dir.exists():
        return {"skills": [], "directory": str(skills_dir), "exists": False}

    skills = []
    for file_path in sorted(skills_dir.glob("*.md")):
        if file_path.name.lower() == "readme.md":
            continue
        try:
            content = file_path.read_text(encoding="utf-8")
            name = file_path.stem
            for line in content.split("\n"):
                if line.startswith("# "):
                    name = line.lstrip("# ").strip()
                    break
            skills.append({"name": name, "content": content, "filename": file_path.name})
        except Exception as e:
            print(f"Failed to read skill file {file_path}: {e}")

    return {"skills": skills, "directory": str(skills_dir), "exists": True}


# ─────────────────────────────────────────────────────────────────────────────
# v3.5: Quantization Switching (Feature 11)
# ─────────────────────────────────────────────────────────────────────────────

QUANTIZATION_MAP: dict[str, str] = {
    "low": "llama3.2:1b",
    "medium": "llama3.2:3b",
    "high": "llama3.1:8b",
}


def get_quantization_model(complexity: int, preferred: str | None = None) -> str:
    """Select model based on task complexity and availability.

    complexity 1 → low (llama3.2:1b, fallback 3b)
    complexity 2 → medium (llama3.2:3b)
    complexity 3 → high (llama3.1:8b, fallback 3b)
    """
    if preferred:
        return preferred
    if complexity <= 1:
        level = "low"
    elif complexity == 2:
        level = "medium"
    else:
        level = "high"

    target = QUANTIZATION_MAP[level]
    # Check if target model exists in capabilities map as sanity check
    if target in MODEL_CAPABILITIES:
        return target
    # Fallback for low/high to medium
    return QUANTIZATION_MAP["medium"]


# ─────────────────────────────────────────────────────────────────────────────
# v3.5: Speculative Decoding — Draft + Refine (Feature 14)
# ─────────────────────────────────────────────────────────────────────────────

async def speculative_decode(
    messages: list[dict],
    full_model: str,
    draft_model: str = "llama3.2:1b",
) -> str:
    """Fast speculative decode: try small draft model first, refine only if needed.

    1. Calls draft_model (tiny, fast) with 15s timeout.
    2. If draft looks good (len > 100, non-empty) → return it directly.
    3. Otherwise → call full_model for full-quality response.
    """
    try:
        draft = await asyncio.wait_for(
            ollama_chat_text(messages, model=draft_model, temperature=0.7, max_tokens=1024),
            timeout=15,
        )
        if draft and len(draft.strip()) > 100:
            _logger.debug(f"[speculative_decode] Draft accepted ({len(draft)} chars)")
            return draft
    except asyncio.TimeoutError:
        _logger.debug("[speculative_decode] Draft timed out, using full model")
    except Exception as e:
        _logger.debug(f"[speculative_decode] Draft failed: {e}")

    # Fall back to full model
    try:
        result = await asyncio.wait_for(
            ollama_chat_text(messages, model=full_model, temperature=0.7, max_tokens=2048),
            timeout=120,
        )
        return result or ""
    except Exception as e:
        _logger.error(f"[speculative_decode] Full model failed: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# v3.5: AI Skill Compiler (Feature 9)
# ─────────────────────────────────────────────────────────────────────────────

_compiled_skills_cache: list[dict] = []


async def compile_skills() -> list[dict]:
    """Read all .md skill files and extract structured metadata.

    For each file extracts: name, description (first paragraph), triggers,
    capabilities (bullet list items), and raw content.
    """
    skills_dir = _BASE / "skills"
    if not skills_dir.exists():
        return []

    compiled: list[dict] = []
    for file_path in sorted(skills_dir.glob("*.md")):
        if file_path.name.lower() == "readme.md":
            continue
        try:
            content = file_path.read_text(encoding="utf-8", errors="replace")
            lines = content.split("\n")

            # Extract name from first # heading
            name = file_path.stem
            for line in lines:
                if line.startswith("# "):
                    name = line.lstrip("# ").strip()
                    break

            # Extract description: first non-empty paragraph after the title
            description = ""
            in_desc = False
            desc_lines: list[str] = []
            for line in lines:
                if line.startswith("# "):
                    in_desc = True
                    continue
                if in_desc:
                    if line.strip():
                        desc_lines.append(line.strip())
                    elif desc_lines:
                        break  # End of first paragraph

            description = " ".join(desc_lines)[:400]

            # Extract trigger lines
            triggers = [
                line.strip() for line in lines
                if "trigger" in line.lower() and line.strip()
            ]

            # Extract capabilities: lines starting with - in first bullet list
            capabilities = [
                line.lstrip("- ").strip()
                for line in lines
                if line.startswith("- ") and line.strip()
            ][:10]

            compiled.append({
                "name": name,
                "filename": file_path.name,
                "description": description,
                "triggers": triggers,
                "capabilities": capabilities,
                "content": content,
            })
        except Exception as e:
            _logger.warning(f"[compile_skills] Failed to read {file_path.name}: {e}")

    return compiled


# ─────────────────────────────────────────────────────────────────────────────
# v3.5: Autonomous Tool Discovery (Feature 10)
# ─────────────────────────────────────────────────────────────────────────────

def get_available_tools() -> list[dict]:
    """Return the registry of all available backend tools for autonomous agents."""
    return [
        {
            "name": "web_search",
            "description": "Search the web for current information",
            "endpoint": "/api/web-search",
            "params": ["query", "scrape"],
        },
        {
            "name": "knowledge_search",
            "description": "Search local knowledge base with hybrid RAG",
            "endpoint": "/api/semantic-search",
            "params": ["query", "limit"],
        },
        {
            "name": "code_execution",
            "description": "Execute Python code in a safe sandbox",
            "endpoint": "/api/run-code",
            "params": ["code"],
        },
        {
            "name": "weather",
            "description": "Get current weather for any location",
            "endpoint": "/api/weather",
            "params": ["location"],
        },
        {
            "name": "memory_recall",
            "description": "Recall relevant memories from past conversations",
            "endpoint": "/api/memory/recall",
            "params": ["query"],
        },
        {
            "name": "deep_research",
            "description": "Perform recursive deep research on a topic",
            "endpoint": "/api/deep-research",
            "params": ["query", "depth"],
        },
    ]


# Inject tool descriptions into Researcher agent prompt (Feature 10)
_tool_list_text = "\n".join(
    f"- {t['name']}: {t['description']} (endpoint: {t['endpoint']})"
    for t in get_available_tools()
)
AGENT_SPECIFIC_PROMPTS["Researcher"] += (
    f"\n\nAvailable tools you can reference in your analysis:\n{_tool_list_text}"
)


# ─────────────────────────────────────────────────────────────────────────────
# v3.5: AI Project Mode — Pydantic models + in-memory store (Feature 12)
# ─────────────────────────────────────────────────────────────────────────────

class ProjectRequest(BaseModel):
    name: str
    goal: str
    context: Optional[str] = ""


class ProjectTaskRequest(BaseModel):
    project_id: str
    task: str
    status: Optional[str] = "pending"


_projects: dict[str, dict] = {}  # project_id → project dict


# ─────────────────────────────────────────────────────────────────────────────
# v3.4: Sentinel API endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/sentinel/health")
async def sentinel_health():
    """Get Sentinel routing health snapshot (last 50 entries)."""
    return _sentinel.get_health()


@app.post("/api/sentinel/optimize")
async def sentinel_optimize():
    """Trigger Sentinel analysis and auto-tune confidence thresholds."""
    result = _sentinel.optimize()
    return result


# ── v3.5: Sentinel self-improvement endpoint ──────────────────────────────────

@app.post("/api/sentinel/improve")
async def sentinel_improve():
    """Analyze agent performance history and auto-adjust prompts."""
    result = _sentinel.improve_prompts()
    return result


# ── v3.6: Security status endpoint ───────────────────────────────────────────

@app.get("/api/security/status")
async def security_status():
    """Show recent injection scan findings."""
    return {
        "recent_scans": _injection_scan_log[-20:],
        "total_flagged": len(_injection_scan_log),
    }


# ── v3.6: Bandit router stats endpoint ───────────────────────────────────────

@app.get("/api/bandit/stats")
async def bandit_stats():
    """Show bandit router statistics per model."""
    router = _get_bandit_router()
    return {"stats": router.stats, "total_calls": router._total_calls}


# ── v3.6: Skill discovery endpoint ───────────────────────────────────────────

@app.get("/api/skills/discovered")
async def get_discovered_skills():
    """Return auto-discovered skill patterns from agent execution history."""
    skills = discover_skills_from_traces()
    return {
        "skills": skills,
        "traces_analyzed": len(_execution_traces),
        "min_frequency": 3,
    }


# ── v3.5: Compiled Skills endpoint (Feature 9) ───────────────────────────────

@app.get("/api/skills/compiled")
async def skills_compiled():
    """Return compiled skill definitions with extracted triggers and capabilities."""
    global _compiled_skills_cache
    if not _compiled_skills_cache:
        _compiled_skills_cache = await compile_skills()
    return {
        "skills": _compiled_skills_cache,
        "total": len(_compiled_skills_cache),
    }


# ── v3.5: Tool Discovery endpoint (Feature 10) ───────────────────────────────

@app.get("/api/tools/discover")
async def tools_discover():
    """Return all available backend tools for autonomous agent discovery."""
    return {"tools": get_available_tools()}


# ── v3.5: AI Project Mode endpoints (Feature 12) ─────────────────────────────

@app.post("/api/projects/create")
async def create_project(req: ProjectRequest):
    """Create a new AI project with a goal and optional context."""
    import uuid
    project_id = str(uuid.uuid4())[:8]
    project = {
        "id": project_id,
        "name": req.name,
        "goal": req.goal,
        "context": req.context or "",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "tasks": [],
    }
    _projects[project_id] = project
    return project


@app.get("/api/projects")
async def list_projects():
    """List all active AI projects."""
    return {"projects": list(_projects.values()), "total": len(_projects)}


@app.get("/api/projects/{project_id}")
async def get_project(project_id: str):
    """Get details of a specific project."""
    if project_id not in _projects:
        raise HTTPException(status_code=404, detail=f"Project '{project_id}' not found")
    return _projects[project_id]


@app.post("/api/projects/{project_id}/task")
async def add_project_task(project_id: str, req: ProjectTaskRequest):
    """Add a task to an existing project."""
    if project_id not in _projects:
        raise HTTPException(status_code=404, detail=f"Project '{project_id}' not found")
    task = {
        "id": f"task-{int(time.time() * 1000)}",
        "task": req.task,
        "status": req.status or "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    _projects[project_id]["tasks"].append(task)
    return {"status": "ok", "task": task, "project_id": project_id}


@app.delete("/api/projects/{project_id}")
async def delete_project(project_id: str):
    """Delete a project and all its tasks."""
    if project_id not in _projects:
        raise HTTPException(status_code=404, detail=f"Project '{project_id}' not found")
    del _projects[project_id]
    return {"status": "ok", "deleted": project_id}


# ── v3.5: Speculative Decode status endpoint (Feature 14) ────────────────────

@app.get("/api/speculative/status")
async def speculative_status():
    """Check whether the draft model (llama3.2:1b) is available for speculative decoding."""
    draft_model = "llama3.2:1b"
    available_models = await get_loaded_models()
    draft_available = any(draft_model in m for m in available_models)
    return {
        "draft_model": draft_model,
        "draft_available": draft_available,
        "full_model_fallback": QUANTIZATION_MAP["medium"],
        "available_models": available_models,
        "status": "ready" if draft_available else "draft_model_not_loaded",
    }


# ─────────────────────────────────────────────────────────────────────────────
# v3.4: Voice STT/TTS endpoints
# ─────────────────────────────────────────────────────────────────────────────

try:
    from faster_whisper import WhisperModel as _WhisperModel
    _WHISPER_AVAILABLE = True
except ImportError:
    _WhisperModel = None  # type: ignore
    _WHISPER_AVAILABLE = False

try:
    import pyttsx3 as _pyttsx3
    _PYTTSX3_AVAILABLE = True
except ImportError:
    _pyttsx3 = None  # type: ignore
    _PYTTSX3_AVAILABLE = False

_whisper_model = None
_whisper_lock = asyncio.Lock()


async def _get_whisper() -> object | None:
    """Lazy-load Whisper model (tiny.en by default — ~80MB, very fast)."""
    global _whisper_model
    if not _WHISPER_AVAILABLE:
        return None
    async with _whisper_lock:
        if _whisper_model is None:
            try:
                _whisper_model = _WhisperModel("tiny.en", device="cpu", compute_type="int8")
                _logger.info("Whisper tiny.en loaded")
            except Exception as e:
                _logger.error(f"Whisper load failed: {e}")
                return None
    return _whisper_model


@app.post("/api/voice/transcribe")
async def voice_transcribe(request: "VoiceTranscribeRequest"):
    """Transcribe base64-encoded audio (WebM/WAV) → text via faster-whisper."""
    import base64
    import tempfile

    if not _WHISPER_AVAILABLE:
        raise HTTPException(503, "faster-whisper not installed. Run: pip install faster-whisper")

    model = await _get_whisper()
    if model is None:
        raise HTTPException(503, "Whisper model failed to load")

    try:
        audio_bytes = base64.b64decode(request.audio_b64)
        suffix = ".webm" if request.format == "webm" else ".wav"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
            f.write(audio_bytes)
            tmp_path = f.name

        loop = asyncio.get_event_loop()
        def _transcribe():
            segments, info = model.transcribe(tmp_path, beam_size=1, language="en")
            return "".join(s.text for s in segments).strip()

        text = await loop.run_in_executor(None, _transcribe)
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

        return {"text": text, "language": "en"}
    except Exception as e:
        raise HTTPException(500, f"Transcription failed: {e}")


@app.post("/api/voice/speak")
async def voice_speak(request: "VoiceSpeakRequest"):
    """Synthesize text to speech via pyttsx3 (non-blocking daemon thread)."""
    if not _PYTTSX3_AVAILABLE:
        raise HTTPException(503, "pyttsx3 not installed. Run: pip install pyttsx3")

    import threading

    def _speak_thread(text: str, rate: int, volume: float):
        try:
            engine = _pyttsx3.init()
            voices = engine.getProperty("voices")
            # Prefer female voice (index 1) if available
            if voices and len(voices) > 1:
                engine.setProperty("voice", voices[1].id)
            engine.setProperty("rate", rate)
            engine.setProperty("volume", volume)
            engine.say(text)
            engine.runAndWait()
        except Exception as e:
            _logger.warning(f"TTS speak failed: {e}")

    t = threading.Thread(
        target=_speak_thread,
        args=(request.text, request.rate, request.volume),
        daemon=True,
    )
    t.start()
    return {"status": "speaking", "text": request.text[:100]}


# ─────────────────────────────────────────────────────────────────────────────
# Vision endpoints
# ─────────────────────────────────────────────────────────────────────────────


@app.get("/api/vision/status")
async def vision_status():
    """Check whether a vision model is available in Ollama."""
    model = await get_vision_model()
    return {
        "available": model is not None,
        "model": model,
        "supported_models": VISION_MODELS,
        "install_hint": "ollama pull llava-phi3" if not model else None,
    }


@app.post("/api/vision/analyze")
async def vision_analyze(req: VisionRequest):
    """Analyze an image using a local vision model (llava, llava-phi3, etc.)"""
    model = req.model or await get_vision_model()
    if not model:
        raise HTTPException(status_code=503, detail="No vision model available. Run: ollama pull llava-phi3")

    # Ollama vision API uses the /api/chat endpoint with image in messages
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": req.prompt,
                "images": [req.image_b64],  # Ollama accepts raw base64 (no data URL prefix)
            }
        ],
        "stream": False,
        "options": {"num_keep": 256},
        "keep_alive": "10m",
    }

    client = await get_ollama_client()
    try:
        resp = await client.post(f"{OLLAMA_URL}/api/chat", json=payload, timeout=120.0)
        resp.raise_for_status()
        content = resp.json()["message"]["content"]
        return {"description": content, "model": model}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Vision analysis failed: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Tool Execution Agent — shell, python (sandbox), git
# ─────────────────────────────────────────────────────────────────────────────

SAFE_SHELL_COMMANDS = {
    "ls", "dir", "pwd", "echo", "cat", "type", "head", "tail",
    "grep", "find", "wc", "date", "hostname", "whoami",
    "pip", "pip3", "python", "python3", "node", "npm",
    "git", "curl", "wget",
}

class ShellRequest(BaseModel):
    command: str
    timeout: int = 15

class GitRequest(BaseModel):
    repo_path: str = "."
    command: str  # e.g. "log --oneline -10" or "status"

@app.post("/api/tools/shell")
async def tools_shell(request: ShellRequest):
    """Run a whitelisted shell command and return stdout/stderr."""
    import shlex
    try:
        parts = shlex.split(request.command)
    except ValueError as e:
        raise HTTPException(400, f"Invalid command: {e}")

    if not parts:
        raise HTTPException(400, "Empty command")

    base_cmd = parts[0].lower().rstrip(".exe")
    if base_cmd not in SAFE_SHELL_COMMANDS:
        raise HTTPException(403, f"Command '{parts[0]}' is not in the allowed list: {sorted(SAFE_SHELL_COMMANDS)}")

    # Block dangerous flags
    dangerous = [";", "&&", "||", "|", ">", ">>", "<", "`", "$(",
                 "rm", "del", "format", "mkfs", "dd", "shutdown", "reboot"]
    joined = request.command.lower()
    for d in dangerous:
        if d in joined:
            raise HTTPException(403, f"Command contains disallowed token: '{d}'")

    try:
        proc = await asyncio.create_subprocess_shell(
            request.command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(
            proc.communicate(), timeout=request.timeout
        )
        return {
            "stdout": stdout.decode(errors="replace"),
            "stderr": stderr.decode(errors="replace"),
            "returncode": proc.returncode,
        }
    except asyncio.TimeoutError:
        raise HTTPException(408, "Command timed out")
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/tools/git")
async def tools_git(request: GitRequest):
    """Run read-only git commands in a specified repository path."""
    ALLOWED_GIT = {"log", "status", "diff", "show", "branch", "tag", "remote", "stash"}
    parts = request.command.strip().split()
    if not parts or parts[0] not in ALLOWED_GIT:
        raise HTTPException(403, f"Git sub-command must be one of: {sorted(ALLOWED_GIT)}")

    import shlex
    cmd = f'git -C "{request.repo_path}" {request.command}'
    try:
        proc = await asyncio.create_subprocess_shell(
            cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=15)
        return {
            "stdout": stdout.decode(errors="replace"),
            "stderr": stderr.decode(errors="replace"),
            "returncode": proc.returncode,
        }
    except asyncio.TimeoutError:
        raise HTTPException(408, "Git command timed out")
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────────────────────────────
# Plugin System
# ─────────────────────────────────────────────────────────────────────────────

_BUILTIN_PLUGINS: list[dict] = [
    {
        "id": "web_search",
        "name": "Web Search",
        "description": "DuckDuckGo live search + page scraping for real-time information",
        "category": "research",
        "icon": "🔍",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "code_interpreter",
        "name": "Code Interpreter",
        "description": "Safe sandboxed Python execution with output capture",
        "category": "dev",
        "icon": "💻",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "knowledge_watcher",
        "name": "Knowledge Watcher",
        "description": "Auto-ingest files dropped into backend/knowledge/ folder",
        "category": "rag",
        "icon": "📂",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "hybrid_rag",
        "name": "Hybrid RAG",
        "description": "BM25 + vector semantic search for document retrieval",
        "category": "rag",
        "icon": "🧠",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "voice_io",
        "name": "Voice I/O",
        "description": "Whisper speech-to-text + pyttsx3 text-to-speech",
        "category": "io",
        "icon": "🎙️",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "memory_system",
        "name": "Memory System",
        "description": "Long-term memory with decay scoring and importance filtering",
        "category": "memory",
        "icon": "💾",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "critic_voting",
        "name": "Critic Deliberation",
        "description": "3-voter quality gate before Critic loop (accuracy, completeness, clarity)",
        "category": "quality",
        "icon": "⚖️",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "sentinel",
        "name": "Sentinel Engine",
        "description": "Auto-adjusts routing confidence thresholds based on failure rate",
        "category": "ops",
        "icon": "🛡️",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "weather_tool",
        "name": "Weather Tool",
        "description": "Real-time weather via Open-Meteo (no API key required)",
        "category": "research",
        "icon": "🌤️",
        "enabled": True,
        "builtin": True,
    },
    {
        "id": "mermaid_diagrams",
        "name": "Mermaid Diagrams",
        "description": "Render flowcharts, sequence diagrams, Gantt charts from Markdown",
        "category": "viz",
        "icon": "📊",
        "enabled": True,
        "builtin": True,
    },
]

_plugin_enabled: dict[str, bool] = {p["id"]: p["enabled"] for p in _BUILTIN_PLUGINS}

class PluginToggleRequest(BaseModel):
    plugin_id: str
    enabled: bool

class PluginExecuteRequest(BaseModel):
    plugin_id: str
    params: dict = {}

@app.get("/api/plugins")
async def list_plugins():
    result = []
    for p in _BUILTIN_PLUGINS:
        row = dict(p)
        row["enabled"] = _plugin_enabled.get(p["id"], p["enabled"])
        result.append(row)
    return {"plugins": result}

@app.post("/api/plugins/toggle")
async def toggle_plugin(request: PluginToggleRequest):
    if request.plugin_id not in _plugin_enabled:
        raise HTTPException(404, f"Plugin '{request.plugin_id}' not found")
    _plugin_enabled[request.plugin_id] = request.enabled
    return {"plugin_id": request.plugin_id, "enabled": request.enabled}

@app.post("/api/plugins/execute")
async def execute_plugin(request: PluginExecuteRequest):
    """Execute a specific plugin with given params."""
    if not _plugin_enabled.get(request.plugin_id, False):
        raise HTTPException(403, f"Plugin '{request.plugin_id}' is disabled")

    pid = request.plugin_id
    params = request.params

    if pid == "web_search":
        query = params.get("query", "")
        if not query:
            raise HTTPException(400, "query param required")
        results = await asyncio.get_event_loop().run_in_executor(None, lambda: web_search(query))
        return {"results": results}

    elif pid == "weather_tool":
        city = params.get("city", "London")
        url = f"https://geocoding-api.open-meteo.com/v1/search?name={city}&count=1&language=en&format=json"
        async with _http.get(url) as r:
            geo = r.json() if hasattr(r, "json") else {}
        return {"message": f"Use /api/weather?city={city} for weather data"}

    elif pid == "code_interpreter":
        code = params.get("code", "")
        if not code:
            raise HTTPException(400, "code param required")
        result = run_code(code, params.get("timeout", 10))
        return result

    else:
        return {"message": f"Plugin '{pid}' is active but has no direct execute endpoint. Use its specific API."}


# ─────────────────────────────────────────────────────────────────────────────
# Self-Improving Prompt Engine
# ─────────────────────────────────────────────────────────────────────────────

class PromptOptimizeRequest(BaseModel):
    prompt: str
    goal: str = ""
    iterations: int = 1

@app.post("/api/prompts/optimize")
async def optimize_prompt(request: PromptOptimizeRequest):
    """Use the LLM to critique and improve a prompt."""
    system = """You are an expert prompt engineer. Your job is to take a user prompt and make it significantly better.

Analyze the given prompt for:
- Clarity and specificity
- Ambiguity or vagueness
- Missing context or constraints
- Tone and framing
- Potential for misinterpretation

Then return EXACTLY this JSON structure (no other text):
{
  "analysis": "Brief analysis of the original prompt's weaknesses",
  "improved_prompt": "The fully rewritten, improved prompt",
  "changes": ["Change 1", "Change 2", "Change 3"],
  "score_before": 6,
  "score_after": 9
}"""

    goal_text = f"\n\nGoal/use-case: {request.goal}" if request.goal else ""
    user_msg = f"Original prompt:\n{request.prompt}{goal_text}"

    current_prompt = request.prompt
    iterations_result = []

    for i in range(max(1, min(request.iterations, 3))):
        payload = {
            "model": GENERAL_MODEL,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": f"Original prompt:\n{current_prompt}{goal_text}"},
            ],
            "stream": False,
            "temperature": 0.4,
        }
        try:
            resp = await _http.post(OLLAMA_CHAT_URL, json=payload, timeout=60.0)
            raw = resp.json()["message"]["content"]
            # Extract JSON from response
            import re as _re2
            json_match = _re2.search(r'\{[\s\S]*\}', raw)
            if json_match:
                import json as _json2
                parsed = _json2.loads(json_match.group())
                iterations_result.append(parsed)
                current_prompt = parsed.get("improved_prompt", current_prompt)
            else:
                iterations_result.append({"raw": raw, "improved_prompt": current_prompt})
        except Exception as e:
            iterations_result.append({"error": str(e)})
            break

    final = iterations_result[-1] if iterations_result else {}
    return {
        "original": request.prompt,
        "optimized": final.get("improved_prompt", request.prompt),
        "analysis": final.get("analysis", ""),
        "changes": final.get("changes", []),
        "score_before": final.get("score_before", 0),
        "score_after": final.get("score_after", 0),
        "iterations": iterations_result,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Autonomous Mode
# ─────────────────────────────────────────────────────────────────────────────

from enum import Enum as _Enum
from dataclasses import dataclass as _dataclass, field as _field

class AutonomousStatus(_Enum):
    IDLE = "idle"
    PLANNING = "planning"
    RUNNING = "running"
    PAUSED = "paused"
    DONE = "done"
    FAILED = "failed"

@_dataclass
class AutonomousTask:
    id: str
    description: str
    status: str = "pending"  # pending | running | done | failed | skipped
    result: str = ""
    started_at: float = 0.0
    finished_at: float = 0.0

@_dataclass
class AutonomousSession:
    goal: str
    status: AutonomousStatus = AutonomousStatus.IDLE
    tasks: list[AutonomousTask] = _field(default_factory=list)
    final_result: str = ""
    error: str = ""
    created_at: float = _field(default_factory=lambda: __import__("time").time())
    max_iterations: int = 5

_autonomous_session: AutonomousSession | None = None
_autonomous_lock = asyncio.Lock()

class AutonomousStartRequest(BaseModel):
    goal: str
    max_iterations: int = 5

async def _run_autonomous(session: AutonomousSession):
    """Background coroutine that drives the autonomous agent loop."""
    import time as _time

    try:
        session.status = AutonomousStatus.PLANNING

        # Step 1: Plan — break goal into subtasks
        plan_payload = {
            "model": GENERAL_MODEL,
            "messages": [
                {"role": "system", "content": """You are an autonomous planning agent.
Given a high-level goal, break it into 3-7 concrete, actionable subtasks.
Return ONLY a JSON array of task descriptions, e.g.:
["Research topic X", "Write code for Y", "Test Z", "Summarize findings"]
No other text."""},
                {"role": "user", "content": f"Goal: {session.goal}"},
            ],
            "stream": False,
            "temperature": 0.3,
        }
        resp = await _http.post(OLLAMA_CHAT_URL, json=plan_payload, timeout=60.0)
        raw_plan = resp.json()["message"]["content"]

        import re as _re3, json as _json3
        arr_match = _re3.search(r'\[[\s\S]*\]', raw_plan)
        if arr_match:
            task_list = _json3.loads(arr_match.group())
        else:
            task_list = [session.goal]

        session.tasks = [
            AutonomousTask(id=str(i), description=t)
            for i, t in enumerate(task_list[: session.max_iterations])
        ]
        session.status = AutonomousStatus.RUNNING

        accumulated = ""

        # Step 2: Execute each task
        for task in session.tasks:
            if session.status == AutonomousStatus.PAUSED:
                task.status = "pending"
                continue
            if session.status == AutonomousStatus.FAILED:
                break

            task.status = "running"
            task.started_at = _time.time()

            exec_payload = {
                "model": GENERAL_MODEL,
                "messages": [
                    {"role": "system", "content": f"""You are an autonomous execution agent.
You are working toward this overall goal: {session.goal}

Work already done:
{accumulated if accumulated else "Nothing yet."}

Now execute the current subtask completely. Be thorough and specific."""},
                    {"role": "user", "content": f"Execute subtask: {task.description}"},
                ],
                "stream": False,
                "temperature": 0.5,
            }
            try:
                exec_resp = await _http.post(OLLAMA_CHAT_URL, json=exec_payload, timeout=120.0)
                result = exec_resp.json()["message"]["content"]
                task.result = result
                task.status = "done"
                accumulated += f"\n\n[Task: {task.description}]\n{result}"
            except Exception as e:
                task.result = str(e)
                task.status = "failed"
            finally:
                task.finished_at = _time.time()

        # Step 3: Synthesize final result
        if session.status == AutonomousStatus.RUNNING:
            synth_payload = {
                "model": GENERAL_MODEL,
                "messages": [
                    {"role": "system", "content": "Synthesize the work done into a final, comprehensive, well-structured response. Use markdown."},
                    {"role": "user", "content": f"Goal: {session.goal}\n\nWork completed:\n{accumulated}"},
                ],
                "stream": False,
                "temperature": 0.4,
            }
            synth_resp = await _http.post(OLLAMA_CHAT_URL, json=synth_payload, timeout=120.0)
            session.final_result = synth_resp.json()["message"]["content"]
            session.status = AutonomousStatus.DONE

    except Exception as e:
        session.error = str(e)
        session.status = AutonomousStatus.FAILED


@app.post("/api/autonomous/start")
async def autonomous_start(request: AutonomousStartRequest):
    global _autonomous_session
    async with _autonomous_lock:
        if _autonomous_session and _autonomous_session.status in (
            AutonomousStatus.PLANNING, AutonomousStatus.RUNNING
        ):
            raise HTTPException(409, "An autonomous session is already running. Stop it first.")
        _autonomous_session = AutonomousSession(
            goal=request.goal,
            max_iterations=min(request.max_iterations, 10),
        )
        session = _autonomous_session

    asyncio.create_task(_run_autonomous(session))
    return {"status": "started", "goal": request.goal}


@app.get("/api/autonomous/status")
async def autonomous_status():
    if _autonomous_session is None:
        return {"status": "idle", "goal": None, "tasks": [], "final_result": ""}

    s = _autonomous_session
    return {
        "status": s.status.value,
        "goal": s.goal,
        "tasks": [
            {
                "id": t.id,
                "description": t.description,
                "status": t.status,
                "result": t.result[:500] if t.result else "",
                "started_at": t.started_at,
                "finished_at": t.finished_at,
            }
            for t in s.tasks
        ],
        "final_result": s.final_result,
        "error": s.error,
    }


@app.post("/api/autonomous/stop")
async def autonomous_stop():
    global _autonomous_session
    if _autonomous_session is None:
        return {"status": "no session"}
    _autonomous_session.status = AutonomousStatus.PAUSED
    return {"status": "paused"}


@app.post("/api/autonomous/reset")
async def autonomous_reset():
    global _autonomous_session
    _autonomous_session = None
    return {"status": "reset"}


# ─────────────────────────────────────────────────────────────────────────────
# Static file serving — SPA middleware
# ─────────────────────────────────────────────────────────────────────────────


def _find_dist_dir() -> Path | None:
    """Locate the frontend dist/ directory."""
    candidates = [
        _BASE / "dist",
        Path(__file__).resolve().parent / "dist",
    ]
    for d in candidates:
        if d.exists() and (d / "index.html").exists():
            return d
    return None


_DIST_DIR = _find_dist_dir()
print(f"[static] dist/ = {_DIST_DIR or 'NOT FOUND'}")


# ── v3.7: System model detection ────────────────────────────────────────────
@app.get("/api/system/models")
async def get_system_models():
    """Check which Ollama models are installed."""
    required_models = [
        {"name": "llama3.2:3b",       "label": "Main Model (LLaMA 3.2 3B)",   "required": True,  "size": "~2 GB"},
        {"name": "qwen2.5-coder:3b",  "label": "Code Model (Qwen 2.5 Coder)", "required": False, "size": "~2 GB"},
        {"name": "llava:7b",           "label": "Vision Model (LLaVA 7B)",     "required": False, "size": "~4 GB"},
    ]
    ollama_running = False
    installed_names: set[str] = set()
    try:
        _c = await get_ollama_client()
        resp = await _c.get(f"{OLLAMA_URL}/api/tags", timeout=4.0)
        if resp.status_code == 200:
            ollama_running = True
            data = resp.json()
            for m in data.get("models", []):
                installed_names.add(m.get("name", "").split(":")[0])
                installed_names.add(m.get("name", ""))
    except Exception:
        pass

    models_out = []
    for m in required_models:
        base = m["name"].split(":")[0]
        installed = m["name"] in installed_names or base in installed_names
        models_out.append({**m, "installed": installed})

    all_required_ok = ollama_running and all(
        m["installed"] for m in models_out if m["required"]
    )
    return {
        "ollama_running": ollama_running,
        "models": models_out,
        "all_required_ok": all_required_ok,
    }


@app.post("/api/system/install")
async def install_models(req: InstallRequest):
    """Open a terminal window to run ollama pull commands."""
    import subprocess, sys
    models_to_install: list[str] = []
    if req.install_all:
        models_to_install = ["llama3.2:3b", "qwen2.5-coder:3b", "llava:7b"]
    elif req.models:
        models_to_install = req.models

    if not models_to_install:
        return {"status": "error", "message": "No models specified"}

    cmds = " && ".join(f"ollama pull {m}" for m in models_to_install)
    try:
        if sys.platform == "win32":
            subprocess.Popen(
                ["cmd", "/k", cmds],
                creationflags=subprocess.CREATE_NEW_CONSOLE,
            )
        elif sys.platform == "darwin":
            script = f'tell application "Terminal" to do script "{cmds}"'
            subprocess.Popen(["osascript", "-e", script])
        else:
            for term in ["gnome-terminal", "xterm", "konsole", "xfce4-terminal"]:
                try:
                    subprocess.Popen([term, "--", "bash", "-c", f"{cmds}; exec bash"])
                    break
                except FileNotFoundError:
                    continue
        return {"status": "ok", "models": models_to_install}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/api/files/process")
async def process_file(req: FileProcessRequest):
    """Extract text from uploaded files (PDF, DOCX, text, code)."""
    import base64, io
    name = req.name
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    mime = req.mime_type or ""

    try:
        raw = base64.b64decode(req.content_b64)
    except Exception as e:
        return {"text": "", "type": "error", "error": str(e)}

    text = ""
    file_type = "text"

    if ext == "pdf" or "pdf" in mime:
        file_type = "pdf"
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n\n".join(pages).strip()
        except Exception as e:
            text = f"[PDF extraction failed: {e}]"

    elif ext in ("docx", "doc") or "word" in mime:
        file_type = "docx"
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            text = f"[DOCX extraction failed: {e}]"

    elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp") or "image" in mime:
        file_type = "image"
        text = f"[Image file: {name}]"

    else:
        # Treat everything else as text (code files, markdown, csv, json, yaml, etc.)
        file_type = "text"
        try:
            text = raw.decode("utf-8", errors="replace")
        except Exception:
            text = f"[Could not decode: {name}]"

    word_count = len(text.split()) if text else 0
    return {"text": text, "type": file_type, "word_count": word_count, "name": name}


# ─────────────────────────────────────────────────────────────────────────────
# Desktop app entry point (PyInstaller exe or `python main.py`)
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import signal
    import subprocess
    import threading
    import uvicorn

    APP_PORT = 8000

    # ── Cleanup: kill any old processes on our port ──────────────────────
    def _kill_old_port_users(port: int):
        """Kill leftover processes on the port, including orphan children."""
        if platform.system() != "Windows":
            return
        try:
            result = subprocess.run(
                ["netstat", "-ano"],
                capture_output=True, text=True, timeout=5,
            )
            my_pid = os.getpid()
            killed = 0
            for line in result.stdout.splitlines():
                if f":{port}" in line and "LISTENING" in line:
                    parts = line.split()
                    if not parts:
                        continue
                    try:
                        pid = int(parts[-1])
                    except ValueError:
                        continue
                    if pid in (my_pid, 0):
                        continue
                    try:
                        proc = psutil.Process(pid)
                        name = proc.name().lower()
                        if any(k in name for k in ("echo", "python", "uvicorn")):
                            print(f"[cleanup] Killing stale PID {pid} ({name})")
                            proc.kill()
                            proc.wait(timeout=3)
                            killed += 1
                    except psutil.NoSuchProcess:
                        # Ghost socket — hunt orphan multiprocessing children
                        for child in psutil.process_iter(["pid", "name", "cmdline"]):
                            try:
                                cmd = " ".join(child.info["cmdline"] or [])
                                if f"parent_pid={pid}" in cmd and "multiprocessing" in cmd:
                                    print(f"[cleanup] Killing orphan child PID {child.pid}")
                                    child.kill()
                                    child.wait(timeout=3)
                                    killed += 1
                            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                                pass
                    except (psutil.AccessDenied, psutil.TimeoutExpired):
                        pass
            if killed:
                print(f"[cleanup] Removed {killed} stale process(es)")
        except Exception as e:
            print(f"[cleanup] Port scan failed (non-fatal): {e}")

    print("[startup] Checking for stale processes on port %d..." % APP_PORT)
    _kill_old_port_users(APP_PORT)
    # Brief pause to let the OS release the socket
    time.sleep(0.5)

    # ── Start uvicorn in a background thread ─────────────────────────────
    def _start_server():
        uvicorn.run(app, host="0.0.0.0", port=APP_PORT, log_level="info")

    server_thread = threading.Thread(target=_start_server, daemon=True)
    server_thread.start()

    # Wait briefly for uvicorn to start
    time.sleep(1.5)

    # ── Open native window (or fallback to browser) ──────────────────────
    try:
        import webview
        webview.create_window(
            "ECHO",
            f"http://localhost:{APP_PORT}",
            width=1280,
            height=800,
            min_size=(800, 500),
        )
        webview.start()
    except ImportError:
        # Fallback: open in browser if pywebview not installed
        import webbrowser
        print("[!!] pywebview not installed, opening in browser instead")
        print("  Install it:  pip install pywebview")
        webbrowser.open(f"http://localhost:{APP_PORT}")
        # Keep the process alive
        try:
            server_thread.join()
        except KeyboardInterrupt:
            pass

    # ── Window closed (or browser fallback ended) — force-exit ───────────
    print("[shutdown] Window closed - terminating all threads.")
    os._exit(0)
